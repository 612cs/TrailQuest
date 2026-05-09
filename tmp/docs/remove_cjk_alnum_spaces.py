from pathlib import Path
import re

from docx import Document


DOCX = Path("DOCS/lunwen/陈胜202201150626论文-第5章代码表格版.docx")

# Remove only spaces between CJK characters and ASCII letters/digits.
# Keep English-internal spaces and code indentation untouched.
PATTERNS = (
    (re.compile(r"(?<=[\u4e00-\u9fff])\s+(?=[A-Za-z0-9])"), ""),
    (re.compile(r"(?<=[A-Za-z0-9])\s+(?=[\u4e00-\u9fff])"), ""),
)


def normalize_text(text: str) -> tuple[str, int]:
    changed = 0
    for pattern, repl in PATTERNS:
        text, count = pattern.subn(repl, text)
        changed += count
    return text, changed


def normalize_run_text(run) -> int:
    text = run.text
    new_text, count = normalize_text(text)
    if count:
        run.text = new_text
    return count


doc = Document(DOCX)
changes = 0

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        changes += normalize_run_text(run)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    changes += normalize_run_text(run)

doc.save(DOCX)
print(f"changes={changes}")
