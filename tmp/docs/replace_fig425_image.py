from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOCX = Path("DOCS/lunwen/陈胜202201150626论文-第5章代码表格版.docx")
PNG = Path("DOCS/architecture-diagrams/exports/png/景观预测与推荐链路协同示意图.png")


def has_image(paragraph):
    return bool(paragraph._p.findall(".//" + qn("w:drawing")) or paragraph._p.findall(".//" + qn("w:pict")))


doc = Document(DOCX)
target_idx = None

for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip() == "图4.25景观预测与推荐链路协同示意图":
        for j in range(i - 1, -1, -1):
            if has_image(doc.paragraphs[j]):
                target_idx = j
                break
        break

if target_idx is None:
    raise RuntimeError("未找到图4.25前的图片段落。")

paragraph = doc.paragraphs[target_idx]
old_width = None
extent = paragraph._p.find(".//" + qn("wp:extent"))
if extent is not None:
    old_width = int(extent.attrib.get("cx", "0"))

paragraph.clear()
run = paragraph.add_run()
if old_width:
    run.add_picture(str(PNG), width=old_width)
else:
    run.add_picture(str(PNG))

doc.save(DOCX)
print(f"replaced paragraph {target_idx}")
