from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn


DOCX = Path("DOCS/lunwen/陈胜202201150626论文-第5章代码表格版.docx")

PATTERNS = (
    (re.compile(r"(?<=[\u4e00-\u9fff])\s+(?=[A-Za-z0-9])"), ""),
    (re.compile(r"(?<=[A-Za-z0-9])\s+(?=[\u4e00-\u9fff])"), ""),
)


def normalize_text(text: str) -> tuple[str, int]:
    total = 0
    for pattern, repl in PATTERNS:
        text, count = pattern.subn(repl, text)
        total += count
    return text, total


def has_image(paragraph) -> bool:
    return bool(
        paragraph._p.findall(".//" + qn("w:drawing"))
        or paragraph._p.findall(".//" + qn("w:pict"))
    )


def should_skip_paragraph(paragraph, paragraph_index: int) -> bool:
    style = paragraph.style.name if paragraph.style is not None else ""
    if style.startswith("toc") or style.startswith("Heading"):
        return True
    # Skip cover/abstract/table-of-contents front matter. Body starts at chapter 1.
    if paragraph_index < 119:
        return True
    if has_image(paragraph):
        return True
    return False


def rewrite_paragraph(paragraph, text: str):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    first = paragraph.runs[0]
    style = {
        "bold": first.bold,
        "italic": first.italic,
        "underline": first.underline,
        "font_name": first.font.name,
        "font_size": first.font.size,
    }
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = style["bold"]
    run.italic = style["italic"]
    run.underline = style["underline"]
    run.font.name = style["font_name"]
    run.font.size = style["font_size"]
    if style["font_name"]:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), style["font_name"])


doc = Document(DOCX)
changes = 0

for i, paragraph in enumerate(doc.paragraphs):
    if should_skip_paragraph(paragraph, i):
        continue
    new_text, count = normalize_text(paragraph.text)
    if count:
        rewrite_paragraph(paragraph, new_text)
        changes += count

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if has_image(paragraph):
                    continue
                new_text, count = normalize_text(paragraph.text)
                if count:
                    rewrite_paragraph(paragraph, new_text)
                    changes += count

doc.save(DOCX)
print(f"changes={changes}")
