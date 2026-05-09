from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


DOCX = Path("DOCS/lunwen/陈胜202201150626论文-第5章代码表格版.docx")
DRAWIOS = [
    Path("DOCS/architecture-diagrams/sources/协同图/AI推荐与环境辅助决策协同架构图.drawio"),
    Path("DOCS/architecture-diagrams/sources/协同图/景观预测与推荐链路协同示意图.drawio"),
]


def set_para_text(paragraph, text):
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)


doc = Document(DOCX)
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text == "图4.24 AI推荐与环境辅助决策协同架构图":
        set_para_text(paragraph, "图4.24 大模型推荐链路示意图")
    elif text == "图4.25景观预测与推荐链路协同示意图":
        set_para_text(paragraph, "图4.25 随机森林景观预测模型示意图")
    elif "AI推荐与环境辅助决策之间的协同关系如图4.24所示" in text:
        set_para_text(
            paragraph,
            text.replace("AI推荐与环境辅助决策之间的协同关系如图4.24所示", "大模型推荐链路如图4.24所示"),
        )
    elif "景观预测与推荐链路的关系如图4.25所示" in text:
        set_para_text(
            paragraph,
            text.replace("景观预测与推荐链路的关系如图4.25所示", "随机森林景观预测模型结构如图4.25所示"),
        )
doc.save(DOCX)


for path in DRAWIOS:
    tree = ET.parse(path)
    root = tree.getroot()
    changed = False
    for diagram in root.findall(".//diagram"):
        name = diagram.attrib.get("name", "")
        if "AI推荐" in name or "大模型" in name:
            diagram.set("name", "大模型推荐链路示意图")
        elif "景观预测" in name or "随机森林" in name:
            diagram.set("name", "随机森林景观预测模型示意图")

    for cell in root.findall(".//mxCell"):
        if cell.attrib.get("edge") == "1" and cell.attrib.get("value", "").strip():
            style = cell.attrib.get("style", "")
            if "labelBackgroundColor=" not in style:
                if style and not style.endswith(";"):
                    style += ";"
                style += "labelBackgroundColor=none;labelBorderColor=none;"
                cell.set("style", style)
                changed = True
    tree.write(path, encoding="utf-8", xml_declaration=False)
    print(path, "changed", changed)
