from html import unescape
from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path("DOCS/lunwen/陈胜202201150626论文-4.3数据库设计修订.docx")
DST = Path("DOCS/lunwen/陈胜202201150626论文-第5章代码表格版.docx")
SVG_DIR = Path("DOCS/thesis/figures/code-tables/chapter5")


def svg_to_code(svg_path: Path) -> str:
    text = svg_path.read_text(encoding="utf-8")
    values = re.findall(r"<text\b[^>]*>(.*?)</text>", text)
    return "\n".join(unescape(v) for v in values)


def has_image(paragraph) -> bool:
    return bool(
        paragraph._p.findall(".//" + qn("w:drawing"))
        or paragraph._p.findall(".//" + qn("w:pict"))
    )


def previous_non_empty(paragraphs, idx):
    for j in range(idx - 1, -1, -1):
        text = paragraphs[j].text.strip()
        if text:
            return text
    return ""


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def insert_code_table_before(doc, paragraph, code: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    set_table_borders(table)

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)

    # Replace the default empty paragraph with code lines inside a single cell.
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)

    lines = code.splitlines()
    for line_idx, line in enumerate(lines):
        if line_idx:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)

    paragraph._p.addprevious(table._tbl)
    # Remove the table from the document end, because we moved its XML node.
    remove_paragraph(paragraph)


shutil.copy2(SRC, DST)
doc = Document(DST)
codes = [svg_to_code(path) for path in sorted(SVG_DIR.glob("ch5_code_*.svg"))]
code_iter = iter(codes)

replaced = 0
original_paragraphs = list(doc.paragraphs)
for idx, paragraph in enumerate(original_paragraphs):
    if not has_image(paragraph):
        continue
    prev = previous_non_empty(original_paragraphs, idx)
    if prev in {"核心代码：", "核心算法展示：", "核心算法："}:
        try:
            code = next(code_iter)
        except StopIteration as exc:
            raise RuntimeError("代码 SVG 数量少于文档中的代码图片数量") from exc
        insert_code_table_before(doc, paragraph, code)
        replaced += 1

remaining_codes = list(code_iter)
if remaining_codes:
    raise RuntimeError(f"还有 {len(remaining_codes)} 份 SVG 代码未插入，请检查匹配规则。")

if replaced != 11:
    raise RuntimeError(f"预期替换 11 张代码图片，实际替换 {replaced} 张。")

doc.save(DST)
print(DST)
print(f"replaced={replaced}")
