from copy import deepcopy
from pathlib import Path

from docx import Document


DOCX = Path("DOCS/lunwen/陈胜202201150626论文-4.3数据库设计修订.docx")

DESCRIPTIONS = {
    "（1）用户表：用于保存平台用户的基础个人信息、登录信息以及治理状态信息，是系统其他业务对象建立用户归属关系的起始表。": "（1）用户表：用于保存平台用户的基础个人信息、登录信息以及治理状态信息，是系统其他业务对象建立用户归属关系的起始表。该表的核心字段包括用户ID、用户名、邮箱、密码哈希、用户角色、账号状态、头像信息、封禁信息、创建时间等字段。",
    "用户画像表": "（2）用户画像表：用于保存用户徒步画像信息，是搜索偏好和智能推荐的重要用户侧数据基础。该表的核心字段包括画像ID、用户ID、徒步经验等级、常走路线类型、负重偏好、创建时间和更新时间等字段。",
    "路线表": "（3）路线表：用于保存路线基础信息、空间位置描述、审核状态和交互统计数据，是平台路线浏览、发布和审核流程的核心表。该表的核心字段包括路线ID、封面图、路线名称、位置描述、起点经纬度、结构化地点、难度、负重类型、行程类型、评分、距离、海拔、耗时、路线介绍、发布者ID、路线状态、审核信息等字段。",
    "路线轨迹表": "（4）路线轨迹表：用于保存经过解析的结构化轨迹数据，是地图显示、天气定位和景观辅助判断的重要空间数据表。该表的核心字段包括轨迹记录ID、路线ID、媒体文件ID、上传用户ID、轨迹文件格式、轨迹GeoJSON、起终点坐标、边界框、总距离、最高海拔、累计爬升、轨迹时长、解析状态等字段。",
    "标签表": "（5）标签表：用于保存平台统一标签，支撑路线筛选、路线分类和路线特征表达。该表的核心字段包括标签ID和标签名称，其中标签名称需要保持唯一，便于路线标签关系和搜索筛选功能复用。",
    "媒体文件表": "（6）媒体文件表：用于统一管理头像、路线封面、路线相册、轨迹文件和评论图片等媒体资源，使文件存储与业务数据保持解耦。该表的核心字段包括媒体文件ID、上传用户ID、存储服务商、Bucket名称、对象Key、访问地址、业务类型、原始文件名、MIME类型、文件大小、图片宽高、文件状态等字段。",
}

REMOVE_TEXTS = {
    "用户画像表用于保存用户徒步画像，是搜索偏好和推荐计算的重要用户侧数据基础。",
    "路线表用于保存路线基础信息、空间位置描述、审核状态和交互统计数据，是平台路线浏览、发布和审核流程的核心表。",
    "路线轨迹表用于保存经过解析的结构化轨迹数据，是地图显示、天气定位和景观辅助判断的重要空间数据表。",
    "标签表用于保存平台统一标签，支撑路线筛选和路线特征表达。",
    "媒体文件表用于统一管理头像、路线封面、路线相册、轨迹文件等媒体资源，使文件存储与业务数据保持解耦。",
}


def set_text_like(paragraph, text, template_paragraph):
    """Replace paragraph text while preserving the existing thesis paragraph styling."""
    p_pr = deepcopy(template_paragraph._p.pPr) if template_paragraph._p.pPr is not None else None
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    if p_pr is not None:
        paragraph._p.append(p_pr)
    run = paragraph.add_run(text)
    if template_paragraph.runs:
        src = template_paragraph.runs[0]
        run.bold = src.bold
        run.italic = src.italic
        run.underline = src.underline
        run.font.name = src.font.name
        run.font.size = src.font.size


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


doc = Document(DOCX)

template = None
for para in doc.paragraphs:
    if para.text.strip().startswith("（1）用户表："):
        template = para
        break

if template is None:
    raise RuntimeError("未找到用户表介绍段落，无法提取段落样式。")

for para in doc.paragraphs:
    text = para.text.strip()
    if text in DESCRIPTIONS:
        set_text_like(para, DESCRIPTIONS[text], template)

for para in list(doc.paragraphs):
    if para.text.strip() in REMOVE_TEXTS:
        remove_paragraph(para)

doc.save(DOCX)
