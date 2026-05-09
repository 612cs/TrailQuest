from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path("DOCS/lunwen/陈胜202201150626论文-第5章代码表格版.docx")

RECOMMENDATION_ALGORITHM = """Input: 用户请求 q，用户画像 u，会话上下文 c
Output: 推荐路线集合 R，推荐理由 E，追问建议 F
1. prompt <- buildPrompt(q, u, c)
2. parsed <- DeepSeek.parse(prompt, function_call_schema)
3. if parsed.intent != trail_recommendation then return generalAnswer(q)
4. P <- {l, t, d, r, s}
5. query <- mapToTrailQuery(P.location, P.time, P.difficulty, P.distance, P.scenery)
6. candidates <- trailService.pageTrails(query, userId)
7. if candidates is empty then candidates <- fallbackSearch(P.location)
8. env <- landscapePredictionService.predict(candidates, P.time)
9. for each trail T_i in candidates do
10.     M_i <- calcBasicMatch(T_i, P)
11.     U_i <- calcProfileMatch(T_i, u)
12.     H_i <- calcPopularity(T_i.likes, T_i.favorites, T_i.rating)
13.     E_i <- calcEnvironmentScore(T_i, env)
14.     score(T_i) <- alpha*M_i + beta*U_i + gamma*H_i + delta*E_i
15.     reason <- buildReason(T_i, P, u, env)
16. end for
17. R <- topK(sortByScore(candidates), 3)
18. E <- buildRouteFacts(R)
19. F <- buildFollowUps(parsed, R)
20. return R, E, F"""

LANDSCAPE_ALGORITHM = """Input: 路线 ID trailId，预测天数 days
Output: 景观预测结果 P，景观等级 L
1. days <- clamp(days, 1, 7)
2. trail <- queryTrail(trailId)
3. location <- resolveStartCoordinate(trail)
4. track <- queryTrailTrack(trailId)
5. weather <- fetchCurrentAndHourlyWeather(location)
6. astronomy <- fetchSunMoonData(location, days)
7. light <- fetchLightPollution(location)
8. X <- engineerFeatures(weather, astronomy, light, track)
9. sunriseSunset <- rulePredictSunriseSunset(X)
10. cloudSeaProb <- average(randomForestCloudSea.trees.predict(X))
11. rimeProb <- average(randomForestRime.trees.predict(X))
12. icicleProb <- average(randomForestIcicle.trees.predict(X))
13. cloudSeaLevel <- classify(cloudSeaProb)
14. rimeLevel <- classify(rimeProb)
15. icicleLevel <- classify(icicleProb)
16. P <- buildPredictionCards(sunriseSunset, cloudSeaProb, rimeProb, icicleProb)
17. L <- buildLevelSummary(cloudSeaLevel, rimeLevel, icicleLevel)
18. return P, L"""


def set_cell_text(cell, text):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(14)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


doc = Document(DOCX)
updated = 0

for table in doc.tables:
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    if "Input: 用户请求 q，用户画像 u，会话上下文 c" in text:
        set_cell_text(table.cell(0, 0), RECOMMENDATION_ALGORITHM)
        updated += 1
    elif "Input: 路线 ID trailId，预测天数 days" in text:
        set_cell_text(table.cell(0, 0), LANDSCAPE_ALGORITHM)
        updated += 1

if updated != 2:
    raise RuntimeError(f"预期更新 2 个算法表格，实际更新 {updated} 个。")

doc.save(DOCX)
print(f"updated={updated}")
