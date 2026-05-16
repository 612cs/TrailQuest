from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path("/Users/sheng/Documents/code/hiking")
SRC = ROOT / "DOCS/lunwen/定稿/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V3.docx"
OUT = ROOT / "DOCS/lunwen/定稿/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V4.docx"


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def copy_ppr(src_para, dst_para):
    if src_para._p.pPr is not None:
        dst_para._p.append(src_para._p.pPr)


def add_text(paragraph, text):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "宋体")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), "24")
    rPr.append(szcs)
    r.append(rPr)
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    paragraph._p.append(r)


def add_citation(paragraph, ref_no):
    # begin
    r1 = OxmlElement("w:r")
    rPr1 = OxmlElement("w:rPr")
    va1 = OxmlElement("w:vertAlign")
    va1.set(qn("w:val"), "superscript")
    rPr1.append(va1)
    r1.append(rPr1)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fld_begin)
    paragraph._p.append(r1)

    # instr
    r2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    va2 = OxmlElement("w:vertAlign")
    va2.set(qn("w:val"), "superscript")
    rPr2.append(va2)
    r2.append(rPr2)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' HYPERLINK \\l "ref_{ref_no}" '
    r2.append(instr)
    paragraph._p.append(r2)

    # separate
    r3 = OxmlElement("w:r")
    rPr3 = OxmlElement("w:rPr")
    va3 = OxmlElement("w:vertAlign")
    va3.set(qn("w:val"), "superscript")
    rPr3.append(va3)
    r3.append(rPr3)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fld_sep)
    paragraph._p.append(r3)

    # visible text
    r4 = OxmlElement("w:r")
    rPr4 = OxmlElement("w:rPr")
    va4 = OxmlElement("w:vertAlign")
    va4.set(qn("w:val"), "superscript")
    rPr4.append(va4)
    r4.append(rPr4)
    t = OxmlElement("w:t")
    t.text = f'[{ref_no}]'
    r4.append(t)
    paragraph._p.append(r4)

    # end
    r5 = OxmlElement("w:r")
    rPr5 = OxmlElement("w:rPr")
    va5 = OxmlElement("w:vertAlign")
    va5.set(qn("w:val"), "superscript")
    rPr5.append(va5)
    r5.append(rPr5)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r5.append(fld_end)
    paragraph._p.append(r5)


def build_para(paragraph, parts, template_para):
    clear_paragraph(paragraph)
    copy_ppr(template_para, paragraph)
    for kind, value in parts:
        if kind == "text":
            add_text(paragraph, value)
        else:
            add_citation(paragraph, value)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    tpl = doc.paragraphs[152]  # normal web paragraph with working fonts/indent

    p146 = [
        ("text", "Siriaraya等人对步行导航中的质量感知路线推荐进行了综述，将景观性、安全性、舒适性、可达性和环境体验纳入路线推荐的评价标准中"),
        ("cite", 1),
        ("text", "。在文化旅游路线生成方面，Aksenov等人提出了面向动态用户画像的三层文化路线框架，将兴趣点选择、行程序列安排和出行路径生成结合起来，用于支持更具情境适应性的文化线路推荐"),
        ("cite", 2),
        ("text", "。Lim围绕游客兴趣偏好、起止点偏好和时间预算约束，研究了个性化旅游行程推荐问题，强调路线推荐不应只考虑景点热度，还要同时兼顾用户兴趣与行程约束"),
        ("cite", 3),
        ("text", "。De Domenico等人从智慧城市视角提出个性化路由策略，强调个人约束与整体交通状态协同优化的重要性"),
        ("cite", 4),
        ("text", "。Kabassi从旅游推荐系统个性化角度总结了用户建模、上下文感知和移动终端适配等关键问题"),
        ("cite", 5),
        ("text", "。随后，Gavalas等人对移动旅游推荐系统进行了系统梳理，并进一步提出eCOMPASS多模态旅游路线规划器，将景点选择、公共交通换乘和多日行程组织纳入统一框架"),
        ("cite", 6),
        ("cite", 7),
        ("text", "。Renjith等人对上下文感知个性化旅游推荐系统的发展脉络、输入数据和评价指标进行了系统综述"),
        ("cite", 8),
        ("text", "，Gasmi等人则利用多目标优化方法生成个性化行程推荐规则，为陌生城市中的路线规划提供了新的建模思路"),
        ("cite", 9),
        ("text", "。"),
    ]

    p147 = [
        ("text", "伴随着大语言模型的发展，国外研究也开始关注语言模型在推理、行动以及工具调用方面的能力。Schick等人在Toolformer中证明了语言模型可以通过自监督方式学习何时调用外部工具以及如何组织参数"),
        ("cite", 10),
        ("text", "。Yao等人在ReAct中进一步提出将推理过程与行动过程交替组织，使模型能够在思考与执行之间形成更紧密的协同"),
        ("cite", 11),
        ("text", "。Patil等人在Gorilla中讨论了大语言模型与大规模API连接的能力，表明模型在结合外部接口后能够更可靠地完成工具调用任务"),
        ("cite", 12),
        ("text", "。"),
    ]

    p149 = [
        ("text", "国外有关环境预测、自然场景辅助决策的研究也给本文的景观预测模块提供了一定的借鉴。Kim等人采用随机森林方法对韩国区域能见度进行了预测，结果表明机器学习可以用来预报低能见度的情况；Lakra和Avishek对雾的形成原因、分类、预测、检测以及影响做了综述，认为湿度、温度、风速、气溶胶等都会影响雾和低能见度天气的产生"),
        ("cite", 13),
        ("cite", 14),
        ("text", "。Parde等用集合预报的方法建立了雾概率预报的决策支持系统，为户外出行中低能见度风险提示提供了思路"),
        ("cite", 15),
        ("text", "。这说明将气象数据与地理环境、机器学习模型相结合，可以为云海、雾凇、冰挂等户外景观的概率预测提供可行的技术手段。"),
    ]

    p157 = [
        ("text", "就总体而言，国内对于智慧旅游平台的创建、前后端分离系统实现、地图服务接入、用户体验提升、社区互动以及推荐算法应用等方面已经形成了较多研究成果"),
        ("cite", 16),
        ("cite", 18),
        ("cite", 26),
        ("text", "。目前大部分研究仍集中于城市旅游、景区服务或者一般性的平台建设上，对户外徒步这种具有复杂地形、多变天气、多样轨迹文件、景观窗口和安全风险等特点的垂直场景关注还相对有限。"),
    ]

    p159 = [
        ("text", "经过对国内外有关文献的查阅可知，路线规划技术、智慧旅游技术、AI推荐技术和环境预测技术都有各自不同的理论基础和工程应用。国外研究主要集中在质量感知导航、旅游路线生成、工具调用型大语言模型以及气象预测模型等方向"),
        ("cite", 1),
        ("cite", 10),
        ("cite", 13),
        ("text", "；国内研究则更多关注智慧文旅平台的建设、前后端分离工程实践、地图服务接入以及用户体验优化等内容"),
        ("cite", 16),
        ("cite", 18),
        ("cite", 26),
        ("text", "。"),
    ]

    p184 = [
        ("text", "交互部分采用DeepSeek大模型，可以处理用户用自然语言提出的各种需求。用户可以自由地描述自己的目的地、出行时间、可接受的程度以及所喜欢的景观等信息，系统再从中抽取地点、时间、难度、偏好等主要参数。大模型对于模糊表达、组合条件和口语化提问具有更强的适用性，而传统关键词检索难以胜任这类任务。Toolformer、ReAct、Gorilla等研究说明，大语言模型不仅能够完成文本理解，还能够在推理过程中组织外部工具调用、完成API连接并返回结构化结果，这为本系统实现自然语言驱动的路线推荐与服务编排提供了技术依据"),
        ("cite", 10),
        ("cite", 11),
        ("cite", 12),
        ("text", "。"),
    ]

    for idx, parts in [(146,p146),(147,p147),(149,p149),(157,p157),(159,p159),(184,p184)]:
        build_para(doc.paragraphs[idx], parts, tpl)

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
