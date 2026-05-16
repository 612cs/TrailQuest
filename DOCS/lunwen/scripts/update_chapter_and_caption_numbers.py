from pathlib import Path
import re
from docx import Document


ROOT = Path("/Users/sheng/Documents/code/hiking")
SRC = ROOT / "DOCS/lunwen/定稿/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本.docx"
OUT = ROOT / "DOCS/lunwen/定稿/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V1.docx"


ENTITY_TITLES = {
    "用户实体图": "用户实体属性图",
    "用户徒步画像实体图": "用户徒步画像实体属性图",
    "路线实体图": "路线实体属性图",
    "路线轨迹实体图": "路线轨迹实体属性图",
    "标签实体图": "标签实体属性图",
    "媒体文件实体图": "媒体文件实体属性图",
    "景观观测标签实体图": "景观观测标签实体属性图",
    "景观特征快照实体图": "景观特征快照实体属性图",
    "AI会话实体图": "AI会话实体属性图",
    "AI消息实体图": "AI消息实体属性图",
    "后台操作日志实体图": "后台操作日志实体属性图",
}


def shift_heading_number(text: str) -> str:
    def repl(m):
        n = int(m.group(1))
        if n >= 4:
            return str(n - 1) + m.group(2)
        return m.group(0)

    return re.sub(r"^([4-7])([ .])", repl, text)


def shift_subheading_number(text: str) -> str:
    def repl(m):
        n = int(m.group(1))
        if n >= 4:
            return str(n - 1) + "."
        return m.group(0)

    return re.sub(r"^([4-7])\.", repl, text)


def shift_figure_table_refs(text: str) -> str:
    # chapter-based refs after merging old chapter 2 into chapter 3
    text = re.sub(r"图3\.", "图2.", text)
    text = re.sub(r"图4\.", "图3.", text)
    text = re.sub(r"图5\.", "图4.", text)
    text = re.sub(r"图6\.", "图5.", text)
    text = re.sub(r"表4\.", "表3.", text)
    text = re.sub(r"表5\.", "表4.", text)
    text = re.sub(r"表6\.", "表5.", text)
    text = text.replace("图4.3-5.5", "图4.3-图4.5")
    text = text.replace("图4.11、5.12", "图4.11、图4.12")
    text = text.replace("图4.13、5.14、5.15", "图4.13、图4.14、图4.15")
    text = text.replace("图5.7所示：", "表5.7所示：")
    text = text.replace("图4.7所示：", "表5.7所示：")
    return text


def transform_caption(text: str) -> str:
    m = re.match(r"^(图|表)(\d+)\.(\d+)\s*(.*)$", text)
    if not m:
        return shift_figure_table_refs(text)
    kind, chapter, seq, title = m.groups()
    chapter = int(chapter)
    if kind == "图":
        if chapter >= 3:
            chapter -= 1
    else:
        if chapter >= 4:
            chapter -= 1
    title = ENTITY_TITLES.get(title, title)
    return f"{kind}{chapter}.{seq}  {title}"


def first_text_run(paragraph):
    for run in paragraph.runs:
        if run.text:
            return run
    return None


def rewrite_paragraph_text_keep_tail(paragraph, new_text: str):
    seen = False
    after_tab = False
    for run in paragraph.runs:
        if run.text == "\t":
            after_tab = True
            continue
        if run.text and not seen:
            run.text = new_text
            seen = True
        elif seen and not after_tab and run.text:
            # clear only the remaining title fragments before the TOC page-number field
            run.text = ""
    if not seen:
        paragraph.add_run(new_text)


def main():
    OUT.write_bytes(SRC.read_bytes())
    doc = Document(str(OUT))

    for p in doc.paragraphs:
        style = p.style.name
        text = p.text
        if not text.strip():
            continue

        if style == "Caption" and (text.startswith("图") or text.startswith("表")):
            new_text = transform_caption(text)
            rewrite_paragraph_text_keep_tail(p, new_text)
            continue

        if style in ("Heading 1", "toc 1"):
            base = text.split("\t")[0] if style.startswith("toc") else text
            new_text = shift_heading_number(base)
            if new_text != text:
                rewrite_paragraph_text_keep_tail(p, new_text)
            continue

        if style in ("Heading 2", "Heading 3", "toc 2", "toc 3"):
            base = text.split("\t")[0] if style.startswith("toc") else text
            new_text = shift_subheading_number(base)
            if new_text != text:
                rewrite_paragraph_text_keep_tail(p, new_text)
            continue

        # Normal prose: update textual references only
        for run in p.runs:
            if not run.text:
                continue
            updated = shift_figure_table_refs(run.text)
            if updated != run.text:
                run.text = updated

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
