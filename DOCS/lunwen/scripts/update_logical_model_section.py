from pathlib import Path

from docx import Document


ROOT = Path("/Users/sheng/Documents/code/hiking")
DOC_PATH = ROOT / "DOCS/lunwen/定稿/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V1.docx"


def reset_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = ""


def add_parts(paragraph, parts):
    reset_paragraph(paragraph)
    first = True
    for text, underline in parts:
        run = paragraph.add_run(text)
        if underline:
            run.underline = True
        first = False


def parts_for_model(entity, fields, pk_fields, ending="；"):
    parts = [(f"{entity}（", False)]
    for i, field in enumerate(fields):
        parts.append((field, field in pk_fields))
        if i != len(fields) - 1:
            parts.append(("，", False))
    parts.append(("）" + ending, False))
    return parts


def main():
    doc = Document(str(DOC_PATH))

    doc.paragraphs[366].text = (
        "根据前文概念模型设计以及实体之间的关系，按照关系模型的转换规则，可以将系统中的核心实体"
        "及其关联结构进一步转化为数据库逻辑结构。结合本系统面向户外路线推荐、地图展示、智能对话、"
        "景观预测以及后台治理等业务需求，设计出如下主要关系模型："
    )

    add_parts(
        doc.paragraphs[367],
        parts_for_model(
            "用户",
            ["用户ID", "用户名", "邮箱", "头像", "用户角色", "账号状态", "创建时间"],
            {"用户ID"},
        ),
    )
    add_parts(
        doc.paragraphs[368],
        parts_for_model(
            "用户徒步画像",
            ["画像ID", "用户ID", "徒步经验等级", "常走路线类型", "负重偏好", "更新时间"],
            {"画像ID"},
        ),
    )
    add_parts(
        doc.paragraphs[369],
        parts_for_model(
            "路线",
            ["路线ID", "路线名称", "位置描述", "难度", "距离", "海拔或爬升", "耗时", "审核状态", "发布者ID"],
            {"路线ID"},
        ),
    )
    add_parts(
        doc.paragraphs[370],
        parts_for_model(
            "路线轨迹",
            ["轨迹记录ID", "路线ID", "原始轨迹媒体文件ID", "轨迹GeoJSON", "起终点坐标", "边界范围", "总距离", "累计爬升", "轨迹时长", "解析状态"],
            {"轨迹记录ID"},
        ),
    )
    add_parts(
        doc.paragraphs[371],
        parts_for_model(
            "标签",
            ["标签ID", "标签名称"],
            {"标签ID"},
        ),
    )
    add_parts(
        doc.paragraphs[372],
        parts_for_model(
            "媒体文件",
            ["媒体文件ID", "上传用户ID", "对象键", "访问地址", "业务类型", "文件大小", "文件状态"],
            {"媒体文件ID"},
        ),
    )
    add_parts(
        doc.paragraphs[373],
        parts_for_model(
            "景观观测标签",
            ["观测标签ID", "路线ID", "用户ID", "观测日期", "景观类型", "观测结果", "置信度"],
            {"观测标签ID"},
        ),
    )
    add_parts(
        doc.paragraphs[374],
        parts_for_model(
            "景观特征快照",
            ["特征快照ID", "路线ID", "预测日期", "景观类型", "特征版本", "特征载荷"],
            {"特征快照ID"},
        ),
    )
    add_parts(
        doc.paragraphs[375],
        parts_for_model(
            "AI会话",
            ["会话ID", "用户ID", "会话标题", "模型名称", "会话状态", "更新时间"],
            {"会话ID"},
        ),
    )
    add_parts(
        doc.paragraphs[376],
        parts_for_model(
            "AI消息",
            ["消息ID", "会话ID", "消息角色", "消息内容", "结构化元数据", "创建时间"],
            {"消息ID"},
        ),
    )
    add_parts(
        doc.paragraphs[377],
        parts_for_model(
            "后台操作日志",
            ["日志ID", "操作类型", "操作对象类型", "操作对象ID", "操作人ID", "结果状态", "创建时间"],
            {"日志ID"},
            ending="。",
        ),
    )

    doc.paragraphs[378].text = (
        "除上述核心实体外，系统中还包括用于维护多对多关系和记录用户互动行为的关系表，主要关系模型为："
    )

    # Insert 3 relation-table paragraphs before next heading
    heading_p = doc.paragraphs[379]._p
    relation_specs = [
        ("路线标签关联", ["路线ID", "标签ID"], {"路线ID", "标签ID"}, "；"),
        ("路线点赞", ["点赞ID", "路线ID", "用户ID", "创建时间"], {"点赞ID"}, "；"),
        ("路线收藏", ["收藏ID", "路线ID", "用户ID", "创建时间"], {"收藏ID"}, "。"),
    ]
    inserted = []
    for entity, fields, pks, ending in relation_specs:
        new_p = heading_p.makeelement(heading_p.tag)
        heading_p.addprevious(new_p)
        inserted.append(new_p)

    # reload doc to get inserted paragraphs in sequence
    doc.save(str(DOC_PATH))
    doc = Document(str(DOC_PATH))

    # find 3.4.3 heading index after insertion
    start = None
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "3.4.3 物理模型设计":
            start = idx
            break
    relation_idxs = [start - 3, start - 2, start - 1]

    for idx, spec in zip(relation_idxs, relation_specs):
        entity, fields, pks, ending = spec
        add_parts(doc.paragraphs[idx], parts_for_model(entity, fields, pks, ending))

    doc.save(str(DOC_PATH))


if __name__ == "__main__":
    main()
