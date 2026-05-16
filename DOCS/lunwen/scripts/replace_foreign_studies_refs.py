from pathlib import Path
import shutil

from docx import Document


ROOT = Path("/Users/sheng/Documents/code/hiking")
SRC = ROOT / "DOCS/lunwen/定稿/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V1.docx"
OUT = ROOT / "DOCS/lunwen/定稿/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V2.docx"


def set_plain_paragraph(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_foreign_paragraph_with_citations(paragraph):
    parts = [
        ("Siriaraya等人对步行导航中的质量感知路线推荐进行了综述，将景观性、安全性、舒适性、可达性和环境体验纳入路线推荐的评价标准中", False),
        ("[1]", True),
        ("。在旅游路线生成与行程规划方面，Garcia等人将个性化旅游路线生成划分为推荐、路线生成和路线定制三个步骤，为将游客限制条件、景点选择与公共交通约束结合起来提供了较早的框架", False),
        ("[2]", True),
        ("。De Domenico等人从智慧城市视角提出个性化路由策略，强调个人约束与整体交通状态协同优化的重要性", False),
        ("[3]", True),
        ("。Kabassi从旅游推荐系统个性化角度总结了用户建模、上下文感知和移动终端适配等关键问题", False),
        ("[4]", True),
        ("。随后，Gavalas等人对移动旅游推荐系统进行了系统梳理，并进一步提出eCOMPASS多模态旅游路线规划器，将景点选择、公共交通换乘和多日行程组织纳入统一框架", False),
        ("[5][6]", True),
        ("。近期研究则更关注行程推荐的效率、上下文约束与个性化规则生成。Halder等人通过个性化POI选择和搜索剪枝提升了行程推荐效率", False),
        ("[7]", True),
        ("，Renjith等人对上下文感知个性化旅游推荐系统的发展脉络、输入数据和评价指标进行了系统综述", False),
        ("[8]", True),
        ("，Gasmi等人则利用多目标优化方法生成个性化行程推荐规则，为陌生城市中的路线规划提供了新的建模思路", False),
        ("[9]", True),
        ("。", False),
    ]

    for run in paragraph.runs:
        run.text = ""

    first = True
    for text, superscript in parts:
        run = paragraph.runs[0] if first and paragraph.runs else paragraph.add_run()
        if first and paragraph.runs:
            run = paragraph.runs[0]
            run.text = text
        else:
            run = paragraph.add_run(text)
        run.font.superscript = superscript
        first = False


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    # 1.3.1 国外研究现状
    set_foreign_paragraph_with_citations(doc.paragraphs[146])

    # replace reference entries [2]-[9]
    replacements = {
        618: "Garcia A, Arbelaitz O, Linaza M T, Vansteenwegen P, Souffriau W. Personalized Tourist Route Generation[C]//Daniel F, Facca F M, eds. Current Trends in Web Engineering. Berlin, Heidelberg: Springer, 2010: 486-497.",
        619: "De Domenico M, Lima A, González M C, et al. Personalized routing for multitudes in smart cities[J]. EPJ Data Science, 2015, 4: 1.",
        620: "Kabassi K. Personalizing recommendations for tourists[J]. Telematics and Informatics, 2010, 27(1): 51-66.",
        621: "Gavalas D, Konstantopoulos C, Mastakas K, et al. Mobile recommender systems in tourism[J]. Journal of Network and Computer Applications, 2014, 39: 319-333.",
        622: "Gavalas D, Kasapakis V, Konstantopoulos C, et al. The eCOMPASS multimodal tourist tour planner[J]. Expert Systems with Applications, 2015, 42(21): 7303-7316.",
        623: "Halder S, Lim K H, Chan J, et al. Efficient itinerary recommendation via personalized POI selection and pruning[J]. Knowledge and Information Systems, 2022, 64: 963-993.",
        624: "Renjith S, Sreekumar A, Jathavedan M. An extensive study on the evolution of context-aware personalized travel recommender systems[J]. Information Processing & Management, 2020, 57(1): 102078.",
        625: "Gasmi I, Soui M, Barhoumi K, et al. Recommendation rules to personalize itineraries for tourists in an unfamiliar city[J]. Applied Soft Computing, 2024, 150: 111084.",
    }
    for idx, text in replacements.items():
        set_plain_paragraph(doc.paragraphs[idx], text)

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
