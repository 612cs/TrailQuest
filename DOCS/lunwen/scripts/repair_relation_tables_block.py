from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table


ROOT = Path("/Users/sheng/Documents/code/hiking")
DOC_PATH = ROOT / "DOCS/lunwen/定稿/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V1.docx"


def make_paragraph_like(template_para, text):
    p = OxmlElement("w:p")
    if template_para._p.pPr is not None:
        p.append(deepcopy(template_para._p.pPr))
    r = OxmlElement("w:r")
    if template_para.runs and template_para.runs[0]._r.rPr is not None:
        r.append(deepcopy(template_para.runs[0]._r.rPr))
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_table_like(template_table, rows_data):
    tbl = deepcopy(template_table._tbl)
    trs = tbl.findall(qn("w:tr"))
    header = trs[0]
    data_row = trs[1]
    for tr in trs[1:]:
        tbl.remove(tr)
    for _ in rows_data:
        tbl.append(deepcopy(data_row))
    wrapper = Table(tbl, None)
    for ridx, row in enumerate(rows_data, start=1):
        for cidx, value in enumerate(row):
            wrapper.rows[ridx].cells[cidx].text = value
    return tbl


def clear_caption_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def main():
    doc = Document(str(DOC_PATH))
    body = doc._body._body

    admin_table = doc.tables[10]._tbl
    redis_heading = next(p for p in doc.paragraphs if p.text.strip() == "3.4.4 Redis数据库设计")._p
    redis_caption = next(p for p in doc.paragraphs if "Redis数据库表" in p.text)

    # Remove whatever currently sits between admin table and Redis heading.
    children = list(body)
    start = children.index(admin_table)
    end = children.index(redis_heading)
    for elem in children[start + 1 : end]:
        body.remove(elem)

    # Prepare templates from nearby stable content
    normal_tpl = doc.paragraphs[408]   # Normal (Web)
    caption_tpl = doc.paragraphs[410]  # Caption
    table_tpl = doc.tables[4]          # small 4-col db table

    blocks = [
        (
            "（12）路线标签关联表用于维护路线与标签之间的多对多关系，是路线筛选、分类和特征表达的基础关联表。核心字段有路线ID、标签ID。",
            "表3.12  路线标签关联表",
            [
                ["trail_id", "BIGINT", "路线ID", "PK, FK"],
                ["tag_id", "BIGINT", "标签ID", "PK, FK"],
            ],
        ),
        (
            "（13）路线点赞表用于记录用户对路线的点赞行为，是路线热度统计和互动行为分析的重要支撑表。核心字段有点赞ID、路线ID、用户ID、创建时间。",
            "表3.13  路线点赞表",
            [
                ["id", "BIGINT", "点赞ID", "PK, AUTO_INCREMENT"],
                ["trail_id", "BIGINT", "路线ID", "FK, NOT NULL"],
                ["user_id", "BIGINT", "用户ID", "FK, NOT NULL"],
                ["created_at", "DATETIME", "创建时间", "DEFAULT"],
            ],
        ),
        (
            "（14）路线收藏表用于保存用户对路线的收藏记录，是个人中心收藏管理和路线偏好统计的重要关联表。核心字段有收藏ID、路线ID、用户ID、创建时间。",
            "表3.14  路线收藏表",
            [
                ["id", "BIGINT", "收藏ID", "PK, AUTO_INCREMENT"],
                ["trail_id", "BIGINT", "路线ID", "FK, NOT NULL"],
                ["user_id", "BIGINT", "用户ID", "FK, NOT NULL"],
                ["created_at", "DATETIME", "创建时间", "DEFAULT"],
            ],
        ),
    ]

    # Insert in visual order.
    for desc, caption, rows in blocks:
        body.insert(body.index(redis_heading), make_paragraph_like(normal_tpl, desc))
        body.insert(body.index(redis_heading), make_paragraph_like(caption_tpl, caption))
        body.insert(body.index(redis_heading), make_table_like(table_tpl, rows))

    clear_caption_text(redis_caption, "表3.15  Redis数据库表")
    doc.save(str(DOC_PATH))


if __name__ == "__main__":
    main()
