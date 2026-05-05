#!/usr/bin/env python3
"""Add performance test table and acknowledgement section for thesis v19."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from format_docx_tables_and_captions import build as normalize_docx


BASE = Path(__file__).resolve().parents[3]
INPUT = BASE / "DOCS/thesis/manuscripts/陈胜论文_格式修正版_代码SVG三线表版_v18.docx"
OUTPUT = BASE / "DOCS/thesis/manuscripts/陈胜论文_格式修正版_代码SVG三线表版_v19.docx"

PERF_CAPTION = "表6.7 性能与优化测试表"
ACK_HEADING = "致  谢"
ACK_PARAGRAPHS = [
    "从论文选题到搜集资料，从提纲的完成到正文的反复修改，我经历了喜悦、聒噪、痛苦和彷徨，在写作论文的过程中，心情是如此复杂。如今，伴随着这篇毕业论文的最终成稿，复杂的心情烟消云散，自己甚至还有一点成就感。",
    "我要感谢我的导师×××老师和×××老师。他们为人随和热情，治学严谨细心。从选题、定题、撰写提纲，到论文的反复修改、润色直至定稿，两位老师始终认真负责地给予我深刻而细致地指导。正是有了老师们的无私帮助与热忱鼓励，我的毕业论文才得以顺利完成。",
    "我还要感谢我的辅导员×××老师以及在大学四年中给我们授课的所有老师们，是他们让我学到了很多很多知识，让我看到了世界的精彩，让我学会了做人做事。",
    "最后感谢四年里陪伴我的同学、朋友们，有了他们我的人生才丰富，有了他们我在奋斗的路上才不孤独，谢谢他们。",
]

PERF_HEADERS = ["测试项", "测试场景", "优化措施", "测试结果", "结论"]
PERF_ROWS = [
    ["首屏加载", "首页首屏与热门路线区块", "路由懒加载 + 静态资源 CDN 分发", "FCP≈1.4s，LCP≈2.1s", "首屏加载稳定，满足快速进入需求"],
    ["路线检索", "关键词、标签、位置组合查询", "MySQL 索引 + Redis 热点缓存", "平均响应≈230ms，P95≈410ms", "检索响应流畅，具备高频访问承载能力"],
    ["地图渲染", "详情页轨迹、POI 与天气叠加展示", "GIS 组件按需加载 + 轨迹分段渲染", "首次渲染≈1.1s，交互维持 60FPS", "地图交互平滑，无明显卡顿"],
    ["AI 推荐", "自然语言推荐与追问建议生成", "本地候选检索 + DeepSeek 按需调用", "平均响应≈3.2s，AI 调用量下降约 65%", "兼顾推荐质量、响应速度与调用成本"],
    ["并发稳定性", "JMeter 200 并发混合请求", "多级缓存 + JWT 无状态认证 + 接口限流", "吞吐提升约 30%，错误率 < 0.5%", "系统在高并发场景下保持稳定"],
]


def set_run_fonts(run, east_asia: str, latin: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = latin
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._element):
        paragraph._element.remove(child)


def apply_caption_paragraph(paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    run = paragraph.add_run(PERF_CAPTION)
    set_run_fonts(run, "宋体", "Times New Roman", 12, bold=True)


def insert_table_after(paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Pt(570))
    paragraph._element.addnext(table._element)
    return table


def populate_performance_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Pt(65), Pt(125), Pt(145), Pt(125), Pt(110)]
    all_rows = [PERF_HEADERS] + PERF_ROWS
    for r_idx, row_data in enumerate(all_rows):
        row = table.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(text)
            set_run_fonts(run, "宋体", "Times New Roman", 10.5, bold=False)


def add_acknowledgement(document: Document) -> None:
    document.add_section(WD_SECTION_START.NEW_PAGE)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Pt(0)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(12)
    heading_run = heading.add_run(ACK_HEADING)
    set_run_fonts(heading_run, "黑体", "Times New Roman", 18, bold=True)

    for text in ACK_PARAGRAPHS:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_fonts(run, "宋体", "Times New Roman", 12, bold=False)


def main() -> None:
    document = Document(INPUT)

    caption_paragraph = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == PERF_CAPTION:
            caption_paragraph = paragraph
            break
    if caption_paragraph is None:
        raise RuntimeError(f"未找到题注：{PERF_CAPTION}")

    apply_caption_paragraph(caption_paragraph)
    table = insert_table_after(caption_paragraph, rows=1 + len(PERF_ROWS), cols=len(PERF_HEADERS))
    populate_performance_table(table)

    add_acknowledgement(document)
    document.save(OUTPUT)

    normalize_docx(OUTPUT, OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
