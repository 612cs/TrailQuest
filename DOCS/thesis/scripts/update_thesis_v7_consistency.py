from __future__ import annotations

import argparse
import shutil
import tempfile
from pathlib import Path
from zipfile import ZipFile

from docx import Document


PARAGRAPH_REPLACEMENTS = {
    "随着全民健康意识提升和智慧文旅建设推进，徒步、登山、露营等户外活动逐渐大众化，用户对路线查询、个性化推荐和环境辅助决策的需求不断提高。针对现有地图工具和旅游平台在户外场景中存在信息分散、推荐能力不足、环境风险提示有限等问题，本文设计并实现了一个基于 Vue 的户外路线智能推荐平台。系统采用前后端分离架构，前端基于 Vue、TypeScript、Tailwind CSS 和 DaisyUI 构建交互界面，后端基于 Spring Boot、MyBatis-Plus、MySQL、Redis 和 Spring Security 实现业务处理、数据管理与权限控制。平台实现了路线浏览、条件筛选、社区交互和智能推荐等功能，并引入 DeepSeek 大语言模型与 Function Calling 机制，支持用户通过自然语言完成路线检索与推荐。同时，系统结合高德地图 API、天气数据和随机森林模型，对云海、日出等景观场景提供概率预测。测试结果表明，平台能够完成户外路线查询、推荐和环境辅助提示等主要功能，对提升用户出行前的信息获取效率具有一定参考价值。":
    "随着全民健康意识提升和智慧文旅建设推进，徒步、登山、露营等户外活动逐渐大众化，用户对路线查询、个性化推荐和环境辅助决策的需求不断提高。针对现有地图工具和旅游平台在户外场景中存在信息分散、推荐能力不足、环境风险提示有限等问题，本文设计并实现了一个基于 Vue 的户外路线智能推荐平台。系统采用前后端分离架构，前端基于 Vue、TypeScript、Tailwind CSS 和 DaisyUI 构建交互界面，后端基于 Spring Boot、MyBatis-Plus、MySQL、Redis 和 Spring Security 实现业务处理、数据管理与权限控制。平台实现了路线浏览、条件筛选、社区交互和智能推荐等功能，并引入 DeepSeek 大语言模型与 Function Calling 机制，支持用户通过自然语言完成路线检索与推荐。同时，系统结合高德地图 API、天气数据和随机森林模型，对云海、雾凇、冰挂、日出日落等景观场景提供概率预测。测试结果表明，平台能够完成户外路线查询、推荐和环境辅助提示等主要功能，对提升用户出行前的信息获取效率具有一定参考价值。",
    "视觉渲染与性能优化：UI 层结合 Tailwind CSS 的原子化类名与 DaisyUI 组件库，实现统一的响应式布局。前端针对照片墙、评论列表等长列表场景引入虚拟列表渲染以控制 DOM 规模，并针对 DeepSeek 大语言模型的返回结果使用 SSE 流式文本接收和动态组件挂载机制，使路线卡片能够随对话结果逐步呈现。":
    "视觉渲染与性能优化：UI 层结合 Tailwind CSS 的原子化类名与 DaisyUI 组件库，实现统一的响应式布局。前端针对照片墙、评论列表等长列表场景引入虚拟列表渲染以控制 DOM 规模，并针对 DeepSeek 大语言模型的返回结果使用 WebSocket 流式消息接收和动态组件挂载机制，使路线卡片能够随对话结果逐步呈现。",
    "天气预测：直观展示目标路线区域未来 48 小时的精细化气象预报，以及基于算法推演的云海、日出等自然景观出现概率。":
    "天气预测：直观展示目标路线区域未来 48 小时的精细化气象预报，以及基于随机森林预测链路生成的云海、雾凇、冰挂、日出日落等自然景观出现概率。",
    "智能回答：通过 SSE 流式输出机制，在对话框中实时解答用户的户外百科提问，并在需要时动态挂载路线卡片等交互组件。":
    "智能回答：通过 WebSocket 持续推送机制，在对话框中实时解答用户的户外百科提问，并在需要时动态挂载路线卡片等交互组件。",
    "云海/日出概率预测模型的数据处理与流程":
    "随机森林景观概率预测模型的数据处理与流程",
    "针对气象预测与景观评价，系统在后端服务中集成随机森林预测器，并结合规则修正输出景观概率、等级和风险说明。随机森林通过集成多棵决策树完成分类与概率估计，适合处理气温、湿度、风速、云量、海拔等多因素共同作用的环境预测任务。Kim 等基于随机森林构建了能见度预测模型，证明机器学习方法可用于低能见度场景预测[10]。":
    "针对气象预测与景观评价，系统在后端服务中集成随机森林预测器，并结合规则修正输出云海、雾凇、冰挂、日出日落等景观的概率、等级和风险说明。随机森林通过集成多棵决策树完成分类与概率估计，适合处理气温、湿度、风速、云量、海拔等多因素共同作用的环境预测任务。Kim 等基于随机森林构建了能见度预测模型，证明机器学习方法可用于低能见度场景预测[10]。",
    "在模型集成的数据处理流中，系统首先根据路线起点或文本位置解析经纬度，再通过天气、天文和光污染相关服务获取当前天气、逐小时预报、日月信息和环境上下文。Lakra 与 Avishek 的综述指出，湿度、边界层结构、地形与凝结核条件是影响雾形成与能见度变化的关键因素[11]；因此，本研究将湿度、云量、风速、露点差、海拔和降水等变量编码为模型输入，用于推断云海、雾凇、冰挂等景观概率。":
    "在模型集成的数据处理流中，系统首先根据路线起点或文本位置解析经纬度，再通过天气、天文和光污染相关服务获取当前天气、逐小时预报、日月信息和环境上下文。Lakra 与 Avishek 的综述指出，湿度、边界层结构、地形与凝结核条件是影响雾形成与能见度变化的关键因素[11]；因此，本研究将湿度、云量、风速、露点差、海拔和降水等变量编码为模型输入，用于推断云海、雾凇、冰挂、日出日落等景观概率。",
    "智能服务方面，大模型接口能够承担自然语言意图解析和推荐话术生成任务，Function Calling 思路适合将用户需求转化为结构化查询参数。景观预测模块在后端服务中集成随机森林预测器和规则修正逻辑，用于输出云海、雾凇、冰挂等景观概率及风险提示。因此，本系统在技术选型和工程实现上具备可行性。":
    "智能服务方面，大模型接口能够承担自然语言意图解析和推荐话术生成任务，Function Calling 思路适合将用户需求转化为结构化查询参数。景观预测模块在后端服务中集成随机森林预测器和规则修正逻辑，用于输出云海、雾凇、冰挂、日出日落等景观概率及风险提示。因此，本系统在技术选型和工程实现上具备可行性。",
    "模块描述：天气与景观预测模块负责路线位置解析、天气查询、天文数据获取、环境特征构造和随机森林景观概率预测，为路线详情和 AI 推荐提供环境辅助决策信息。":
    "模块描述：天气与景观预测模块负责路线位置解析、天气查询、天文数据获取、环境特征构造以及云海、雾凇、冰挂、日出日落四类景观的随机森林概率预测，为路线详情和 AI 推荐提供环境辅助决策信息。",
    "该流程图展示了景观预测的核心过程：系统以路线坐标和轨迹海拔为基础，拉取气象、天文和光污染数据，完成特征工程后输入随机森林预测器，输出不同景观的概率、等级和风险提示。":
    "该流程图展示了景观预测的核心过程：系统以路线坐标和轨迹海拔为基础，拉取气象、天文和光污染数据，完成特征工程后输入随机森林预测器，输出云海、雾凇、冰挂、日出日落等景观的概率、等级和风险提示。",
    "后端 LandscapePredictionServiceImpl 将路线环境上下文统一传入不同预测器，云海、雾凇和冰挂等景观分别由随机森林预测器给出概率与说明。":
    "后端 LandscapePredictionServiceImpl 将路线环境上下文统一传入景观预测链路，云海、雾凇、冰挂和日出日落等景观由随机森林预测器与规则修正逻辑共同输出概率与说明。",
    "在随机森林景观预测服务的验证中，系统选取历史天气样例构造输入特征，并对模型输出的景观概率、风险等级和提示文本进行检查。由于当前数据集规模有限，本文仅将该模型作为辅助决策模块进行验证，重点考察其输入解析、特征构造和结果返回流程是否完整。后续若接入更长时间跨度的实测样本，可进一步使用准确率、召回率、AUC 或相关系数等指标评价预测效果。":
    "在随机森林景观预测服务的验证中，系统选取历史天气样例构造输入特征，并对模型输出的景观概率、风险等级和提示文本进行检查。本文已基于随机森林实现云海、雾凇、冰挂和日出日落概率预测；受历史样本规模限制，本阶段验证重点放在输入解析、特征构造、概率输出和结果展示链路是否完整。后续若接入更长时间跨度的实测样本，可进一步使用准确率、召回率、AUC 或相关系数等指标评价预测效果。",
    "本文围绕户外路线发现、路线发布、社区互动、后台审核和智能推荐等需求，设计并实现了一个基于 Vue 3 与 Spring Boot 的户外路线智能推荐平台。系统采用前后端分离架构，前端负责路线浏览、地图展示、轨迹上传和 AI 对话交互，后端负责用户认证、路线管理、评论互动、媒体存储和审核日志等业务处理。在智能服务方面，系统将 DeepSeek 大模型用于自然语言意图解析与推荐话术生成，将随机森林模型用于天气与景观预测辅助，从而形成“人找路线”与条件化推荐相结合的服务方式。测试结果表明，系统主要功能能够按预期完成，整体实现满足毕业设计阶段的应用验证要求。":
    "本文围绕户外路线发现、路线发布、社区互动、后台审核和智能推荐等需求，设计并实现了一个基于 Vue 3 与 Spring Boot 的户外路线智能推荐平台。系统采用前后端分离架构，前端负责路线浏览、地图展示、轨迹上传和 AI 对话交互，后端负责用户认证、路线管理、评论互动、媒体存储和审核日志等业务处理。在智能服务方面，系统将 DeepSeek 大模型用于自然语言意图解析与推荐话术生成，将随机森林模型用于云海、雾凇、冰挂和日出日落概率预测与环境辅助判断，从而形成“人找路线”与条件化推荐相结合的服务方式。测试结果表明，系统主要功能能够按预期完成，整体实现满足毕业设计阶段的应用验证要求。",
}


SVG_REPLACEMENTS = {
    "进入 SSE / 对话通道": "进入 WebSocket 对话通道",
    "SSE / 对话通道": "WebSocket 对话通道",
    "SSE 流式文本接收": "WebSocket 流式消息接收",
    "通过 SSE 流式输出机制": "通过 WebSocket 持续推送机制",
}


def set_paragraph_text(paragraph, text: str) -> bool:
    if paragraph.text != text and paragraph.text in PARAGRAPH_REPLACEMENTS:
        if paragraph.runs:
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = text
        else:
            paragraph.add_run(text)
        return True
    return False


def update_document(docx_path: Path) -> int:
    doc = Document(docx_path)
    changed = 0

    for paragraph in doc.paragraphs:
        old_text = paragraph.text
        new_text = PARAGRAPH_REPLACEMENTS.get(old_text)
        if new_text is not None and set_paragraph_text(paragraph, new_text):
            changed += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    old_text = paragraph.text
                    new_text = PARAGRAPH_REPLACEMENTS.get(old_text)
                    if new_text is not None and set_paragraph_text(paragraph, new_text):
                        changed += 1

    doc.save(docx_path)
    return changed


def update_embedded_svgs(docx_path: Path) -> int:
    temp_dir = Path(tempfile.mkdtemp(prefix="docx_v7_"))
    try:
        with ZipFile(docx_path) as archive:
            archive.extractall(temp_dir)

        changed = 0
        media_dir = temp_dir / "word" / "media"
        for svg_path in media_dir.glob("*.svg"):
            raw = svg_path.read_text(encoding="utf-8", errors="ignore")
            updated = raw
            for old, new in SVG_REPLACEMENTS.items():
                updated = updated.replace(old, new)
            if updated != raw:
                svg_path.write_text(updated, encoding="utf-8")
                changed += 1

        rebuilt = temp_dir / "rebuilt.docx"
        with ZipFile(rebuilt, "w") as archive:
            for path in sorted(temp_dir.rglob("*")):
                if path == rebuilt or path.is_dir():
                    continue
                archive.write(path, path.relative_to(temp_dir))

        shutil.move(rebuilt, docx_path)
        return changed
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()

    paragraph_changes = update_document(args.docx)
    svg_changes = update_embedded_svgs(args.docx)
    print(f"paragraph_changes={paragraph_changes}")
    print(f"svg_changes={svg_changes}")


if __name__ == "__main__":
    main()
