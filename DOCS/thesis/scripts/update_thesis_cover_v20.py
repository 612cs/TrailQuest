#!/usr/bin/env python3
"""Rebuild thesis cover metadata block with a borderless table layout."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


BASE = Path(__file__).resolve().parents[3]
INPUT = BASE / "DOCS/thesis/manuscripts/陈胜论文_格式修正版_代码SVG三线表版_v20.docx"
OUTPUT = BASE / "DOCS/thesis/manuscripts/陈胜论文_格式修正版_代码SVG三线表版_v22.docx"

ROWS = [
    ("论文题目：", "基于Vue的户外路线智能推荐平台的设计与实现"),
    ("学生姓名：", "陈 胜"),
    ("学    号：", "202201150240"),
    ("二级学院：", "计算机科学与工程学院"),
    ("专    业：", "软件工程"),
    ("班    级：", "2022软件工程六班"),
    ("指导教师：", "昌明权"),
]


def set_run_fonts(run, east_asia: str, latin: str, size_pt: float) -> None:
    run.font.name = latin
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_cell_border(cell, edge: str, val: str, size: str | None = None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    el = tc_borders.find(qn(f"w:{edge}"))
    if el is None:
        el = OxmlElement(f"w:{edge}")
        tc_borders.append(el)
    el.set(qn("w:val"), val)
    if size is not None:
        el.set(qn("w:sz"), size)
    elif el.get(qn("w:sz")) is not None:
        del el.attrib[qn("w:sz")]
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "000000")


def set_cell_margins(cell, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = tc_mar.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def clear_cell(cell) -> None:
    for p in cell.paragraphs:
        p.clear()


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "FFFFFF")


def main() -> None:
    document = Document(INPUT)
    if not document.tables:
        raise RuntimeError("未找到封面信息表格")

    table = document.tables[0]
    if len(table.rows) != len(ROWS) or len(table.columns) != 2:
        raise RuntimeError("封面信息表格结构与预期不一致")

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    for row, (label, value) in zip(table.rows, ROWS):
        row.height = Pt(34)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        label_cell = row.cells[0]
        value_cell = row.cells[1]
        label_cell.width = Pt(135)
        value_cell.width = Pt(345)
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        set_cell_margins(label_cell, top=0, bottom=8, left=0, right=0)
        set_cell_margins(value_cell, top=0, bottom=8, left=0, right=0)

        clear_cell(label_cell)
        clear_cell(value_cell)

        lp = label_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lp.paragraph_format.first_line_indent = Pt(0)
        lp.paragraph_format.left_indent = Pt(0)
        lp.paragraph_format.right_indent = Pt(0)
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        lp.paragraph_format.line_spacing = 1.0
        lrun = lp.add_run(label)
        set_run_fonts(lrun, "宋体", "Times New Roman", 15)

        vp = value_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vp.paragraph_format.first_line_indent = Pt(0)
        vp.paragraph_format.left_indent = Pt(0)
        vp.paragraph_format.right_indent = Pt(0)
        vp.paragraph_format.space_before = Pt(0)
        vp.paragraph_format.space_after = Pt(0)
        vp.paragraph_format.line_spacing = 1.0
        vrun = vp.add_run(value)
        set_run_fonts(vrun, "宋体", "Times New Roman", 14)

        for edge in ("top", "left", "right"):
            set_cell_border(value_cell, edge, "nil")
        set_cell_border(value_cell, "bottom", "single", "4")
        for edge in ("top", "left", "right", "bottom"):
            set_cell_border(label_cell, edge, "nil")

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
