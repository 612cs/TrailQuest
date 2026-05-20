#!/usr/bin/env python3
"""Expand the domestic and foreign research status section in the thesis."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


BASE = Path(__file__).resolve().parents[3]
DOCX_PATH = BASE / "DOCS/thesis/manuscripts/final/陈胜论文.docx"


FOREIGN_PARAGRAPHS = [
    "国外在路线规划、空间信息服务与个性化推荐领域起步较早，研究路径已经从经典最短路径计算逐步发展到面向场景质量、用户偏好与动态约束的综合推荐。Siriaraya 等围绕步行导航提出质量感知路线导航综述，指出景观性、安全性、舒适性、可达性和环境体验已经成为路线推荐的重要评价维度[1]。在旅游场景中，Zhang 从大数据视角讨论智能旅游路线推荐方法，Liu 与 Wang 则分别尝试将知识图谱、知识映射等结构化方法引入个性化旅游路线推荐，使旅游路线推荐从单一条件匹配转向知识组织、兴趣建模和路线语义关联[2-4]。",
    "随着大语言模型的发展，国外研究开始关注自然语言交互与工具调用能力在推荐系统中的作用。Sen 等提出 Simple Action Model，讨论了大语言模型顺序调用外部工具完成复杂任务的可行性[8]；Wang 等提出 ToolFlow，通过合成自然连贯的工具调用对话提升模型在多步骤任务中的调用能力[9]；Qin 等进一步研究通用大语言模型开放域 Function Calling 能力的释放方式[10]。这些研究表明，大模型不只能够生成文本，还可以作为意图理解、任务拆解和外部服务编排的中枢，为路线推荐系统从“关键词搜索”升级为“自然语言需求理解 + 工具链执行”提供了重要方法依据。",
    "在智能推荐接受机制方面，Li 等从 Quantified Self 视角研究消费者对 AI 推荐的接受过程，指出用户画像、自我量化数据和推荐解释会影响用户对 AI 推荐结果的信任与采纳[26]。该研究启示本文在设计智能推荐模块时，不应只关注推荐结果是否命中条件，还应关注推荐理由、环境风险说明和追问建议等解释性信息。对于户外路线推荐而言，用户往往需要在时间、体力、天气、景观和安全之间进行权衡，因此系统需要将推荐结果组织成可理解、可比较、可追问的交互形式。",
    "在环境预测与自然场景辅助决策方面，国外研究也为本文的景观预测模块提供了参考。Kim 等基于随机森林方法开展韩国区域能见度预测研究，证明机器学习模型能够利用气象要素对低能见度场景进行预测[11]；Lakra 和 Avishek 对雾形成因素、分类、预测、检测及影响进行了综述，指出湿度、温度、风速、气溶胶等因素会共同影响雾和低能见度天气的形成[12]。这些成果说明，将气象数据、地理环境和机器学习模型结合起来，可以为云海、雾凇、冰挂等户外景观概率预测提供可行技术路径。",
    "总体来看，国外研究已经形成了较为完整的技术谱系：一方面，路线推荐从最短路计算扩展到质量感知、知识图谱和用户偏好建模；另一方面，大语言模型工具调用和随机森林等机器学习方法又为复杂需求理解、外部数据接入和环境预测提供了新的支撑。但这些成果多分散在导航推荐、旅游推荐、AI 工具调用和气象预测等不同领域，直接面向中文户外徒步路线场景，并同时整合 GIS 地图、天气数据、用户画像、AI 对话和景观预测的系统研究仍相对有限。",
]

DOMESTIC_PARAGRAPHS = [
    "国内在智慧旅游、地图服务与个性化推荐方面发展迅速，已经形成了以地图平台、内容平台与旅游服务平台协同演进的应用生态。符琳蓉、汪明峰从服务生态系统视角总结智慧旅游价值共创研究进展，指出智慧旅游正在从单一信息展示转向平台协同、数据连接与多主体价值共创[5,20]。周景兰则从互联网时代智慧旅游网络平台服务模式出发，强调智慧旅游平台需要在资源整合、在线服务和用户体验之间形成持续优化机制[22]。这些研究说明，国内智慧文旅研究已经不再停留于景点信息数字化，而是更加关注平台化服务、用户参与和智能化运营。",
    "在系统工程实现方面，国内大量研究采用 Spring Boot、Vue 等技术栈构建旅游、消防、绩效考核等业务系统。胡金宇基于 Spring Boot 和 Vue 框架完成企业绩效考核系统设计与实现，说明前后端分离架构在权限管理、数据维护和业务流程管理方面具有较强适用性[13]；童学洲基于 Spring Boot 框架实现物联网智慧消防系统，体现了 Spring Boot 在设备接入、状态监控和后台管理场景中的工程支撑能力[14]。在旅游平台方向，肖程鸣、曾志颖设计实现了基于 Spring Boot 和 Vue 的红色智慧旅游平台，黄镓升、邓舒婷则完成了基于 Spring Boot 的南宁旅游 APP 设计与实现[21,25]。这些成果为本文采用 Vue + Spring Boot 构建户外路线推荐平台提供了直接的工程参考。",
    "在智慧文旅交互与用户体验方面，国内研究开始重视多源数据、可视化表达和用户体验优化。刘文榕围绕数字孪生技术研究智慧文旅泛在交互设计，强调通过数字空间映射和多端交互提升文旅场景的信息感知能力[15]；孙彦捷以苏南传统民居建筑视觉图谱为对象开展交互设计研究，体现了视觉图谱在文化资源组织和交互呈现中的价值[16]；薛慧从用户体验角度研究智慧文旅平台优化，指出平台需要在信息架构、交互流程和服务内容上贴合用户真实使用场景[17]。这些研究对本文前端路线展示、地图交互、社区内容瀑布流和个人画像页面设计具有启发意义。",
    "在智能推荐与地图服务方面，国内研究也逐步关注算法透明性、用户行为和地理信息表达。王娜、何文静、孙倬围绕大学生对智能推荐算法的应对行为展开研究，说明用户并非被动接受算法结果，而会根据自身目标、平台机制和推荐可信度主动调整使用行为[19]。乔振华、王珺祺基于百度地图完成吉林省智慧旅游地图设计与实现，表明地图 API 与旅游资源数据结合可以有效提升旅游信息展示和路线引导能力[24]。Wei C. 关于旅游景点图像识别与深度学习推荐的研究，则从图像识别和智能推荐角度说明视觉内容也可以成为旅游推荐的重要数据来源[27]。这些研究共同说明，推荐系统需要同时考虑用户行为、地理空间数据和内容表达方式。",
    "在算法支撑方面，国内对于随机森林等机器学习方法的应用研究也较为活跃。马玉彬等将贝叶斯优化随机森林用于储层预测，说明随机森林模型在复杂特征组合、非线性关系拟合和特征重要性分析方面具有较强适应性[23]。虽然该研究场景并非旅游领域，但其方法思想对本文构建天气与景观预测算法具有借鉴意义。结合陈慕花关于计算机技术在智慧旅游中的应用总结可以看出，云计算、数据分析、智能算法和可视化交互正在共同推动智慧旅游服务从信息展示走向辅助决策[18]。",
    "总体来看，国内研究在智慧旅游平台建设、前后端分离系统实现、地图服务接入、用户体验优化和推荐算法应用方面已经积累了较丰富成果[13-27]。但现有研究大多聚焦城市旅游、景区服务或通用平台建设，对户外徒步这一具有复杂地形、动态天气、轨迹文件、景观窗口和安全风险特征的垂直场景关注仍然不足。尤其是在将自然语言推荐、地图轨迹渲染、天气查询、景观概率预测和社区内容治理整合为一个完整系统方面，仍存在进一步研究和实现空间。",
]

SUMMARY_PARAGRAPHS = [
    "综合国内外研究现状可以看出，路线规划、智慧旅游、AI 推荐和环境预测技术已经具备较为坚实的理论与工程基础[1-27]。国外研究更强调质量感知路线导航、大语言模型工具调用和气象预测模型，国内研究则在智慧文旅平台建设、前后端分离实现、地图服务接入和用户体验优化方面形成了较多实践成果。",
    "然而，现有研究与应用仍然存在三个方面的不足。其一，许多系统的核心设计仍偏重城市道路、标准景区或普通旅游场景，面对高山、峡谷、森林等自然户外环境时，对复杂地形、轨迹解析、动态天气、景观窗口和安全风险等关键因素支持不够充分。其二，部分平台虽然引入了智能推荐思想，但仍主要依赖标签匹配、热度排序或静态条件筛选，对用户自然语言需求、体力偏好、负重方式和出行目的之间的关联理解不足。其三，AI 对话、GIS 地图、天气服务、机器学习预测和社区内容治理往往分散在不同系统或不同研究主题中，缺少面向户外出行决策全过程的协同设计。",
    "因此，围绕户外路线推荐这一具体场景，研究一套融合前端交互、后端业务、地图能力、大语言模型和景观预测的综合平台，既能够回应现实中的使用痛点，也能够推动相关技术在垂直领域中的落地应用。本课题正是在上述研究基础上展开，目标是设计并实现一套更贴近户外出行决策过程的智能路线推荐平台，使用户能够在同一系统中完成路线浏览、条件筛选、AI 咨询、天气查看、景观判断、路线发布和社区互动等操作。",
]

REFERENCE_FIXES = {
    "乔振华,王祺.基于百度地图的吉林省智慧旅游地图的设计与实现[J].无线互联科技,2025,22(15):97-100+109.": "乔振华,王珺祺.基于百度地图的吉林省智慧旅游地图的设计与实现[J].无线互联科技,2025,22(15):97-100+109.",
    "WeiC.Tourist attraction image recognition and intelligent recommendation based on deep learning[J].Journal of Computational Methods in Sciences and Engineering,2025,25(4):3066-3079.": "Wei C. Tourist attraction image recognition and intelligent recommendation based on deep learning[J]. Journal of Computational Methods in Sciences and Engineering, 2025, 25(4): 3066-3079.",
}


def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def clone_paragraph_after(anchor: Paragraph, text: str, template: Paragraph) -> Paragraph:
    new_p = deepcopy(template._p)
    for child in list(new_p):
        if child.tag.endswith("}r"):
            new_p.remove(child)
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    set_paragraph_text(paragraph, text)
    return paragraph


def remove_between(start_heading: Paragraph, end_heading: Paragraph) -> None:
    body = start_heading._p.getparent()
    current = start_heading._p.getnext()
    while current is not None and current is not end_heading._p:
        nxt = current.getnext()
        body.remove(current)
        current = nxt


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph_text(paragraph) == text:
            return paragraph
    raise RuntimeError(f"未找到段落：{text}")


def insert_section_after_heading(heading: Paragraph, next_heading: Paragraph, texts: list[str]) -> None:
    template_element = heading._p.getnext()
    if template_element is None:
        raise RuntimeError(f"标题后缺少正文模板：{paragraph_text(heading)}")
    template = Paragraph(template_element, heading._parent)
    remove_between(heading, next_heading)
    anchor = heading
    for text in texts:
        anchor = clone_paragraph_after(anchor, text, template)


def main() -> None:
    document = Document(DOCX_PATH)
    foreign_heading = find_paragraph(document, "1.3.1 国外研究现状")
    domestic_heading = find_paragraph(document, "1.3.2 国内研究现状")
    summary_heading = find_paragraph(document, "1.3.3 现状综述")
    next_heading = find_paragraph(document, "1.4 本文的主要工作与论文结构")

    insert_section_after_heading(foreign_heading, domestic_heading, FOREIGN_PARAGRAPHS)
    domestic_heading = find_paragraph(document, "1.3.2 国内研究现状")
    summary_heading = find_paragraph(document, "1.3.3 现状综述")
    insert_section_after_heading(domestic_heading, summary_heading, DOMESTIC_PARAGRAPHS)
    summary_heading = find_paragraph(document, "1.3.3 现状综述")
    next_heading = find_paragraph(document, "1.4 本文的主要工作与论文结构")
    insert_section_after_heading(summary_heading, next_heading, SUMMARY_PARAGRAPHS)

    for paragraph in document.paragraphs:
        text = paragraph_text(paragraph)
        if text in REFERENCE_FIXES:
            set_paragraph_text(paragraph, REFERENCE_FIXES[text])

    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
