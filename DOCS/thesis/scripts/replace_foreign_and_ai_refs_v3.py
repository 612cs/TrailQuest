from pathlib import Path
import shutil

from docx import Document


ROOT = Path("/Users/sheng/Documents/code/hiking")
SRC = ROOT / "DOCS/thesis/manuscripts/revisions/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V2.docx"
OUT = ROOT / "DOCS/thesis/manuscripts/revisions/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V3.docx"


def set_runs(paragraph, parts):
    for run in paragraph.runs:
        run.text = ""
    first = True
    for text, superscript in parts:
        if first and paragraph.runs:
            run = paragraph.runs[0]
            run.text = text
        else:
            run = paragraph.add_run(text)
        run.font.superscript = superscript
        first = False


def set_plain(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    # 1.3.1 paragraph with genuine foreign route-recommendation papers
    set_runs(
        doc.paragraphs[146],
        [
            ("Siriaraya等人对步行导航中的质量感知路线推荐进行了综述，将景观性、安全性、舒适性、可达性和环境体验纳入路线推荐的评价标准中", False),
            ("[1]", True),
            ("。在文化旅游路线生成方面，Aksenov等人提出了面向动态用户画像的三层文化路线框架，将兴趣点选择、行程序列安排和出行路径生成结合起来，用于支持更具情境适应性的文化线路推荐", False),
            ("[2]", True),
            ("。Lim围绕游客兴趣偏好、起止点偏好和时间预算约束，研究了个性化旅游行程推荐问题，强调路线推荐不应只考虑景点热度，还要同时兼顾用户兴趣与行程约束", False),
            ("[3]", True),
            ("。De Domenico等人从智慧城市视角提出个性化路由策略，强调个人约束与整体交通状态协同优化的重要性", False),
            ("[4]", True),
            ("。Kabassi从旅游推荐系统个性化角度总结了用户建模、上下文感知和移动终端适配等关键问题", False),
            ("[5]", True),
            ("。随后，Gavalas等人对移动旅游推荐系统进行了系统梳理，并进一步提出eCOMPASS多模态旅游路线规划器，将景点选择、公共交通换乘和多日行程组织纳入统一框架", False),
            ("[6][7]", True),
            ("。Renjith等人对上下文感知个性化旅游推荐系统的发展脉络、输入数据和评价指标进行了系统综述", False),
            ("[8]", True),
            ("，Gasmi等人则利用多目标优化方法生成个性化行程推荐规则，为陌生城市中的路线规划提供了新的建模思路", False),
            ("[9]", True),
            ("。", False),
        ],
    )

    set_runs(
        doc.paragraphs[147],
        [
            ("伴随着大语言模型的发展，国外研究也开始关注语言模型在推理、行动以及工具调用方面的能力。Schick等人在Toolformer中证明了语言模型可以通过自监督方式学习何时调用外部工具以及如何组织参数", False),
            ("[10]", True),
            ("。Yao等人在ReAct中进一步提出将推理过程与行动过程交替组织，使模型能够在思考与执行之间形成更紧密的协同", False),
            ("[11]", True),
            ("。Patil等人在Gorilla中讨论了大语言模型与大规模API连接的能力，表明模型在结合外部接口后能够更可靠地完成工具调用任务", False),
            ("[12]", True),
            ("。", False),
        ],
    )

    set_plain(
        doc.paragraphs[148],
        "这类研究说明，大语言模型已经不再只是简单的文本生成工具，而能够在意图识别、步骤推理、工具调用以及结果解释之间建立协同机制。对于户外路线推荐系统而言，这种能力意味着系统可以更自然地理解用户需求，并把自然语言请求转化为可执行的结构化检索与服务调用过程。"
    )

    # 2.1.3 AI and GIS technology paragraph
    set_runs(
        doc.paragraphs[184],
        [
            ("交互部分采用DeepSeek大模型，可以处理用户用自然语言提出的各种需求。用户可以自由地描述自己的目的地、出行时间、可接受的程度以及所喜欢的景观等信息，系统再从中抽取地点、时间、难度、偏好等主要参数。大模型对于模糊表达、组合条件和口语化提问具有更强的适用性，而传统关键词检索难以胜任这类任务。Toolformer、ReAct、Gorilla等研究说明，大语言模型不仅能够完成文本理解，还能够在推理过程中组织外部工具调用、完成API连接并返回结构化结果，这为本系统实现自然语言驱动的路线推荐与服务编排提供了技术依据", False),
            ("[10][11][12]", True),
            ("。", False),
        ],
    )

    # Replace refs [2]-[12]
    replacements = {
        618: "Aksenov P, Kemperman A D A M, Arentze T A. Toward personalised and dynamic cultural routing: a three-level approach[J]. Procedia Environmental Sciences, 2014, 22: 257-269.",
        619: "Lim K H. Personalized Recommendation of Travel Itineraries based on Tourist Interests and Preferences[C]//Adjunct Proceedings of the 24th Conference on User Modeling, Adaptation and Personalization. Aachen: CEUR-WS.org, 2016.",
        620: "De Domenico M, Lima A, González M C, et al. Personalized routing for multitudes in smart cities[J]. EPJ Data Science, 2015, 4: 1.",
        621: "Kabassi K. Personalizing recommendations for tourists[J]. Telematics and Informatics, 2010, 27(1): 51-66.",
        622: "Gavalas D, Konstantopoulos C, Mastakas K, et al. Mobile recommender systems in tourism[J]. Journal of Network and Computer Applications, 2014, 39: 319-333.",
        623: "Gavalas D, Kasapakis V, Konstantopoulos C, et al. The eCOMPASS multimodal tourist tour planner[J]. Expert Systems with Applications, 2015, 42(21): 7303-7316.",
        624: "Renjith S, Sreekumar A, Jathavedan M. An extensive study on the evolution of context-aware personalized travel recommender systems[J]. Information Processing & Management, 2020, 57(1): 102078.",
        625: "Gasmi I, Soui M, Barhoumi K, et al. Recommendation rules to personalize itineraries for tourists in an unfamiliar city[J]. Applied Soft Computing, 2024, 150: 111084.",
        626: "Schick T, Dwivedi-Yu J, Dessì R, et al. Toolformer: Language Models Can Teach Themselves to Use Tools[J]. Advances in Neural Information Processing Systems, 2023, 36: 68539-68551.",
        627: "Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[EB/OL]. arXiv, 2022: arXiv:2210.03629.",
        628: "Patil S G, Zhang T, Wang X, et al. Gorilla: Large Language Model Connected with Massive APIs[EB/OL]. arXiv, 2023: arXiv:2305.15334.",
    }
    for idx, text in replacements.items():
        set_plain(doc.paragraphs[idx], text)

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
