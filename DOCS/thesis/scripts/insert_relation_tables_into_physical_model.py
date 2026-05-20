from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/sheng/Documents/code/hiking")
DOC_PATH = ROOT / "DOCS/thesis/manuscripts/revisions/基于Vue的户外路线智能推荐平台的设计与实现陈胜_副本-章节更新V1.docx"


def set_paragraph_text(paragraph: Paragraph, text: str):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def build_table_from_template(template_tbl, rows_data):
    tbl = deepcopy(template_tbl)
    trs = tbl.findall(qn("w:tr"))
    # keep header only
    for tr in trs[1:]:
        tbl.remove(tr)
    base_row = trs[1]
    for _ in rows_data:
        tbl.append(deepcopy(base_row))
    return tbl


def set_table_texts(table: Table, rows_data):
    # header stays unchanged
    for r_idx, row_data in enumerate(rows_data, start=1):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value


def main():
    doc = Document(str(DOC_PATH))

    # Existing template pieces
    normal_tpl = deepcopy(doc.paragraphs[408]._p)
    caption_tpl = deepcopy(doc.paragraphs[410]._p)
    table_tpl = deepcopy(doc.tables[4]._tbl)  # 标签表，4列，样式与三线表一致

    # Insert before 3.4.4 heading
    heading = doc.paragraphs[411]._p

    blocks = [
        (
            "（12）路线标签关联表用于维护路线与标签之间的多对多关系，是路线筛选、分类与特征表达的基础关联表。主要字段有路线ID、标签ID。",
            "表3.12  路线标签关联表",
            [
                ["trail_id", "BIGINT", "路线ID", "PK, FK"],
                ["tag_id", "BIGINT", "标签ID", "PK, FK"],
            ],
        ),
        (
            "（13）路线点赞表用于记录用户对路线的点赞行为，是路线热度统计和互动行为分析的主要支撑表。核心字段有点赞ID、路线ID、用户ID、创建时间。",
            "表3.13  路线点赞表",
            [
                ["id", "BIGINT", "点赞ID", "PK, AUTO_INCREMENT"],
                ["trail_id", "BIGINT", "路线ID", "FK, NOT NULL"],
                ["user_id", "BIGINT", "用户ID", "FK, NOT NULL"],
                ["created_at", "DATETIME", "创建时间", "DEFAULT"],
            ],
        ),
        (
            "（14）路线收藏表用于保存用户对路线的收藏记录，是路线偏好统计和个人中心收藏管理的重要关联表。核心字段有收藏ID、路线ID、用户ID、创建时间。",
            "表3.14  路线收藏表",
            [
                ["id", "BIGINT", "收藏ID", "PK, AUTO_INCREMENT"],
                ["trail_id", "BIGINT", "路线ID", "FK, NOT NULL"],
                ["user_id", "BIGINT", "用户ID", "FK, NOT NULL"],
                ["created_at", "DATETIME", "创建时间", "DEFAULT"],
            ],
        ),
    ]

    inserted_tbls = []
    for desc, caption, rows in reversed(blocks):
        tbl_xml = build_table_from_template(table_tpl, rows)
        cap_p = deepcopy(caption_tpl)
        desc_p = deepcopy(normal_tpl)
        heading.addprevious(tbl_xml)
        heading.addprevious(cap_p)
        heading.addprevious(desc_p)
        inserted_tbls.append((desc_p, cap_p, tbl_xml, desc, caption, rows))

    # update redis caption sequence after insertion
    set_paragraph_text(doc.paragraphs[414], "表3.15  Redis数据库表")

    doc.save(str(DOC_PATH))

    # Reload to wrap new elements and populate texts safely
    doc = Document(str(DOC_PATH))
    # locate inserted descriptions/captions by text order around heading
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "3.4.4 Redis数据库设计":
            start = i
            break
    if start is None:
        raise RuntimeError("Cannot find 3.4.4 heading after insertion")

    insertion_start = start - 6  # desc/caption pairs and intervening table anchors
    # Fill nearby desc/captions
    desc_indices = [start - 6, start - 4, start - 2]
    cap_indices = [start - 5, start - 3, start - 1]
    for (desc, caption, rows), di, ci in zip(
        [(b[3], b[4], b[5]) for b in inserted_tbls[::-1]], desc_indices, cap_indices
    ):
        set_paragraph_text(doc.paragraphs[di], desc)
        set_paragraph_text(doc.paragraphs[ci], caption)

    # Tables are inserted before Redis table, so become tables[11], [12], [13]; Redis shifts to [14]
    for ti, (_, _, rows) in zip([11, 12, 13], [(b[3], b[4], b[5]) for b in inserted_tbls[::-1]]):
        set_table_texts(doc.tables[ti], rows)

    doc.save(str(DOC_PATH))


if __name__ == "__main__":
    main()
