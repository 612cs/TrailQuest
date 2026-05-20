#!/usr/bin/env python3
"""Clean thesis references and convert citations to superscript hyperlinks."""

from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


BASE = Path(__file__).resolve().parents[3]
DEFAULT_DOCX_PATH = BASE / "DOCS/thesis/manuscripts/final/陈胜论文.docx"

FINAL_REFS = [
    "Siriaraya P, Wang Y, Zhang Y, et al. Beyond the Shortest Route: A Survey on Quality-Aware Route Navigation for Pedestrians[J]. IEEE Access, 2020, 8: 135569-135590.",
    "Zhang L. Intelligent tourism route recommendation method based on big data[J]. International Journal of Autonomous and Adaptive Communications Systems, 2020, 13(4): 329-341.",
    "Liu X. Design of personalized tourism route recommendation system based on knowledge graph[C]//2022 International Conference on Intelligent Transportation, Big Data & Smart City. IEEE, 2022: 102-106.",
    "Wang C. Personalised leisure tourism route recommendation method based on knowledge map[J]. International Journal of Reasoning-based Intelligent Systems, 2024, 16(1): 37-42.",
    "符琳蓉, 汪明峰. 智慧旅游价值共创研究进展: 基于服务生态系统视角[J]. 福建师范大学学报(自然科学版), 2025, 41(05): 31-38.",
    "肖程鸣, 曾志颖. 基于Spring Boot和Vue的红色智慧旅游平台设计与实现[J]. 软件, 2022, 43(07): 30-33,38.",
    "陈慕花. 计算机技术在智慧旅游中的应用研究[J]. 旅游与摄影, 2025(11): 31-33.",
    "Wang Z, Zeng X, Liu W, et al. ToolFlow: Boosting LLM Tool-Calling Through Natural and Coherent Dialogue Synthesis[C]//Proceedings of NAACL 2025 Long Papers. ACL, 2025: 4246-4263.",
    "Qin S, Zhu Y, Mu L, et al. Meta-Tool: Unleash Open-World Function Calling Capabilities of General-Purpose Large Language Models[C]//Proceedings of the 63rd Annual Meeting of the Association for Computational Linguistics. ACL, 2025: 30653-30677.",
    "Kim B Y, Cha J W, Chang K H, et al. Visibility Prediction over South Korea Based on Random Forest[J]. Atmosphere, 2021, 12(5): 552.",
    "Lakra K, Avishek K. A review on factors influencing fog formation, classification, forecasting, detection and impacts[J]. Rendiconti Lincei, 2022, 33(2): 319-353.",
    "王娜, 何文静, 孙倬. 行随“算”迁或逆“算”而行: 大学生智能推荐算法应对行为研究[J]. 情报理论与实践, 2025, 48(12): 127-136.",
    "黄镓升, 邓舒婷. 基于Spring Boot的南宁旅游APP的设计与实现[A]. 全国高等学校计算机教育研究会. 第32届计算机新科技与教育学术会议论文集[C]. 南宁: 南宁学院信息工程学院, 2025: 187-191.",
    "周景兰. 互联网时代智慧旅游网络平台服务模式构建策略[J]. 太原城市职业技术学院学报, 2025(08): 27-29.",
    "马玉彬, 刘仕友, 曹丹平. 基于贝叶斯优化的随机森林在储层预测的应用[J]. 地球物理学报, 2025, 68(08): 3247-3257.",
    "乔振华, 王珺祺. 基于百度地图的吉林省智慧旅游地图的设计与实现[J]. 无线互联科技, 2025, 22(15): 97-100,109.",
    "Li A, Ding Z, Sun C, et al. Recommending AI based on Quantified Self: Investigating the mechanism of consumer acceptance of AI recommendations[J]. Electronic Markets, 2024, 34(1): 57.",
    "Wei C. Tourist attraction image recognition and intelligent recommendation based on deep learning[J]. Journal of Computational Methods in Sciences and Engineering, 2025, 25(4): 3066-3079.",
    "Su X, He J, Ren J, et al. Personalized Chinese Tourism Recommendation Algorithm Based on Knowledge Graph[J]. Applied Sciences, 2022, 12(20): 10226.",
    "Zheng X Y, Han B T, Ni Z. Tourism Route Recommendation Based on A Multi-Objective Evolutionary Algorithm Using Two-Stage Decomposition and Pareto Layering[J]. IEEE/CAA Journal of Automatica Sinica, 2023, 10(2): 486-500.",
    "Parde A N, Ghude S D, Dhangar N G, et al. Operational Probabilistic Fog Prediction Based on Ensemble Forecast System: A Decision Support System for Fog[J]. Atmosphere, 2022, 13(10): 1608.",
    "曹浩, 黎杰, 谢彬. 基于SpringBoot+Vue的桂林龙胜各族自治县的旅游信息系统设计[J]. 现代信息科技, 2024, 8(16): 102-106.",
    "李晟曈, 刘哲, 俞定国, 等. 基于Vue和SpringBoot的乡村文旅平台设计与实现[J]. 现代计算机, 2023, 29(08): 98-103.",
    "戴亚哲, 李尤, 赵利宏, 等. 基于SpringBoot+Vue的文旅平台设计与研究[J]. 无线互联科技, 2024, 21(21): 70-72.",
]

FOREIGN_PARAGRAPHS = [
    "国外在路线规划、空间信息服务与个性化推荐领域起步较早，研究路径已经从经典最短路径计算逐步发展到面向场景质量、用户偏好与动态约束的综合推荐。Siriaraya 等围绕步行导航提出质量感知路线导航综述，指出景观性、安全性、舒适性、可达性和环境体验已经成为路线推荐的重要评价维度[1]。在旅游场景中，Zhang 从大数据视角讨论智能旅游路线推荐方法，Liu 与 Wang 则分别尝试将知识图谱、知识映射等结构化方法引入个性化旅游路线推荐[2-4]。近年研究进一步强化了多源数据和算法模型在旅游推荐中的作用，Su 等提出基于知识图谱的中文旅游个性化推荐算法，Zheng 等从多目标演化算法角度研究旅游路线推荐，为个性化路线生成和多约束排序提供了新的方法参考[19-20]。",
    "随着大语言模型的发展，国外研究开始关注自然语言交互与工具调用能力在推荐系统中的作用。Wang 等提出 ToolFlow，通过合成自然连贯的工具调用对话提升模型在多步骤任务中的调用能力[8]；Qin 等进一步研究通用大语言模型开放域 Function Calling 能力的释放方式[9]。这些研究表明，大模型不只能够生成文本，还可以作为意图理解、任务拆解和外部服务编排的中枢，为路线推荐系统从“关键词搜索”升级为“自然语言需求理解 + 工具链执行”提供了重要方法依据。",
    "在智能推荐接受机制方面，Li 等从 Quantified Self 视角研究消费者对 AI 推荐的接受过程，指出用户画像、自我量化数据和推荐解释会影响用户对 AI 推荐结果的信任与采纳[17]。该研究启示本文在设计智能推荐模块时，不应只关注推荐结果是否命中条件，还应关注推荐理由、环境风险说明和追问建议等解释性信息。对于户外路线推荐而言，用户往往需要在时间、体力、天气、景观和安全之间进行权衡，因此系统需要将推荐结果组织成可理解、可比较、可追问的交互形式。",
    "在环境预测与自然场景辅助决策方面，国外研究也为本文的景观预测模块提供了参考。Kim 等基于随机森林方法开展韩国区域能见度预测研究，证明机器学习模型能够利用气象要素对低能见度场景进行预测[10]；Lakra 和 Avishek 对雾形成因素、分类、预测、检测及影响进行了综述，指出湿度、温度、风速、气溶胶等因素会共同影响雾和低能见度天气的形成[11]。Parde 等进一步从集合预报角度构建雾概率预报决策支持系统，为户外出行中的低能见度风险提示提供了方法参考[21]。这些成果说明，将气象数据、地理环境和机器学习模型结合起来，可以为云海、雾凇、冰挂等户外景观概率预测提供可行技术路径。",
    "总体来看，国外研究已经形成了较为完整的技术谱系：一方面，路线推荐从最短路计算扩展到质量感知、知识图谱和多目标优化；另一方面，大语言模型工具调用和随机森林等机器学习方法又为复杂需求理解、外部数据接入和环境预测提供了新的支撑。但这些成果多分散在导航推荐、旅游推荐、AI 工具调用和气象预测等不同领域，直接面向中文户外徒步路线场景，并同时整合 GIS 地图、天气数据、用户画像、AI 对话和景观预测的系统研究仍相对有限。",
]

DOMESTIC_PARAGRAPHS = [
    "国内在智慧旅游、地图服务与个性化推荐方面发展迅速，已经形成了以地图平台、内容平台与旅游服务平台协同演进的应用生态。符琳蓉、汪明峰从服务生态系统视角总结智慧旅游价值共创研究进展，指出智慧旅游正在从单一信息展示转向平台协同、数据连接与多主体价值共创[5]。周景兰则从互联网时代智慧旅游网络平台服务模式出发，强调智慧旅游平台需要在资源整合、在线服务和用户体验之间形成持续优化机制[14]。这些研究说明，国内智慧文旅研究已经不再停留于景点信息数字化，而是更加关注平台化服务、用户参与和智能化运营。",
    "在系统工程实现方面，国内大量研究采用 Spring Boot、Vue 等技术栈构建旅游服务系统。肖程鸣、曾志颖设计实现了基于 Spring Boot 和 Vue 的红色智慧旅游平台[6]，黄镓升、邓舒婷完成了基于 Spring Boot 的南宁旅游 APP 设计与实现[13]，曹浩等设计了基于 SpringBoot+Vue 的桂林龙胜各族自治县旅游信息系统，并在系统中融合可视化旅游数据、在线地图、旅游攻略和路线规划等功能[22]。这些成果说明，前后端分离架构在旅游资源展示、用户管理、路线信息维护、社区互动和后台运营管理方面具有较强适用性，也为本文采用 Vue + Spring Boot 构建户外路线推荐平台提供了直接工程参考。",
    "在乡村文旅与路线推荐平台方面，李晟曈等实现了基于 Vue 和 SpringBoot 的乡村文旅平台，通过二维码标识、景点详情、精确搜索和乡村文旅信息推荐实现多主体信息集成[23]；戴亚哲等则在 SpringBoot+Vue 文旅平台中引入粒子群优化算法实现旅游景点最佳路线推荐[24]。这些研究与本文的路线展示、地图交互、推荐排序和后台管理等功能具有较强相似性，也说明国内近年来已经开始将路线推荐算法与文旅平台工程实现结合起来。",
    "在智慧文旅平台优化方面，陈慕花从计算机技术应用角度总结了智慧旅游场景中数据处理、可视化和智能服务的支撑作用[7]。在地图服务方向，乔振华、王珺祺基于百度地图完成吉林省智慧旅游地图设计与实现，表明地图 API 与旅游资源数据结合可以有效提升旅游信息展示和路线引导能力[16]。这些研究对本文前端路线展示、地图交互、轨迹渲染和天气位置解析等功能具有启发意义。",
    "在智能推荐与算法应用方面，国内研究也逐步关注算法透明性、用户行为和模型方法。王娜、何文静、孙倬围绕大学生对智能推荐算法的应对行为展开研究，说明用户并非被动接受算法结果，而会根据自身目标、平台机制和推荐可信度主动调整使用行为[12]。马玉彬等将贝叶斯优化随机森林用于储层预测，体现了随机森林模型在复杂特征组合、非线性关系拟合和特征重要性分析方面的适应性[15]。Wei C. 关于旅游景点图像识别与深度学习推荐的研究，则从图像识别和智能推荐角度说明视觉内容也可以成为旅游推荐的重要数据来源[18]。",
    "总体来看，国内研究在智慧旅游平台建设、前后端分离系统实现、地图服务接入、用户体验优化、社区互动和推荐算法应用方面已经积累了较丰富成果[5-7,12-18,22-24]。但现有研究大多聚焦城市旅游、景区服务或通用平台建设，对户外徒步这一具有复杂地形、动态天气、轨迹文件、景观窗口和安全风险特征的垂直场景关注仍然不足。尤其是在将自然语言推荐、地图轨迹渲染、天气查询、景观概率预测和社区内容治理整合为一个完整系统方面，仍存在进一步研究和实现空间。",
]

SUMMARY_PARAGRAPHS = [
    "综合国内外研究现状可以看出，路线规划、智慧旅游、AI 推荐和环境预测技术已经具备较为坚实的理论与工程基础。国外研究更强调质量感知路线导航、旅游推荐算法、大语言模型工具调用和气象预测模型[1-4,8-11,17,19-21]，国内研究则在智慧文旅平台建设、前后端分离实现、地图服务接入、社区互动、路线推荐和用户体验优化方面形成了较多实践成果[5-7,12-18,22-24]。",
    "然而，将这些研究成果直接应用到中文户外徒步场景时，仍会遇到若干限制。许多系统以城市道路、标准景区或普通旅游服务为主要对象，对高山、峡谷、森林等自然环境下的轨迹解析、动态天气、景观窗口和安全风险支持不足。部分平台虽然引入了智能推荐思想，但推荐逻辑仍以标签匹配、热度排序或静态条件筛选为主，对用户自然语言需求、体力偏好、负重方式和出行目的之间的关联理解不够充分。同时，AI 对话、GIS 地图、天气服务、景观预测和社区内容治理通常分散在不同系统或研究主题中，尚未形成面向户外出行决策全过程的协同实现。",
    "基于上述分析，本文将研究范围收束到户外路线推荐这一具体应用场景，重点关注路线数据组织、地图轨迹展示、自然语言推荐、天气与景观辅助判断以及社区互动治理之间的衔接关系。系统目标不是单独实现某一种算法，而是在一个可运行的平台中验证多类技术如何共同服务于用户的路线选择与出行准备。",
]

TEXT_REPLACEMENTS = {
    "“AI与GIS”": "“AI 与 GIS”",
    "LLM虽然": "LLM 虽然",
    "近期关于工具调用与函数调用链的研究也表明，LLM 在结构化任务编排和外部工具协同方面具备较强扩展潜力[8-10]。": "近期关于工具调用与函数调用的研究也表明，LLM 在结构化任务编排和外部工具协同方面具备较强扩展潜力[8-9]。",
    "Sen 等提出的顺序函数调用链、Wang 等关于 ToolFlow 的研究以及 Qin 等关于开放域函数调用能力的工作，都说明工具调用已成为大模型走向工程化系统的重要路径[8-10]。": "Wang 等关于 ToolFlow 的研究以及 Qin 等关于开放域函数调用能力的工作，都说明工具调用已成为大模型走向工程化系统的重要路径[8-9]。",
    "该思路与顺序函数调用、自然对话驱动工具调用以及开放域工具选择等研究形成了较好的方法呼应[8-10]。": "该思路与自然对话驱动工具调用以及开放域工具选择等研究形成了较好的方法呼应[8-9]。",
    "Kim 等基于随机森林构建了能见度预测模型，并证明其在误差控制和相关性指标上优于传统业务预报结果[11]。": "Kim 等基于随机森林构建了能见度预测模型，并证明其在误差控制和相关性指标上优于传统业务预报结果[10]。",
    "Lakra 与 Avishek 的综述指出，湿度、边界层结构、地形与凝结核条件是影响雾形成与能见度变化的关键因素[12]；": "Lakra 与 Avishek 的综述指出，湿度、边界层结构、地形与凝结核条件是影响雾形成与能见度变化的关键因素[11]；",
}

CITATION_RE = re.compile(r"\[[0-9,\-]+\]")


def ptext(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def clear_and_set(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def clone_after(anchor: Paragraph, text: str, template: Paragraph) -> Paragraph:
    new_p = deepcopy(template._p)
    for child in list(new_p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            new_p.remove(child)
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    clear_and_set(paragraph, text)
    return paragraph


def remove_between(start: Paragraph, end: Paragraph) -> None:
    body = start._p.getparent()
    current = start._p.getnext()
    while current is not None and current is not end._p:
        nxt = current.getnext()
        body.remove(current)
        current = nxt


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if ptext(paragraph) == text:
            return paragraph
    raise RuntimeError(f"未找到段落：{text}")


def replace_section(heading: Paragraph, next_heading: Paragraph, texts: list[str]) -> None:
    template = Paragraph(heading._p.getnext(), heading._parent)
    remove_between(heading, next_heading)
    anchor = heading
    for text in texts:
        anchor = clone_after(anchor, text, template)


def add_bookmark(paragraph: Paragraph, number: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(1000 + number))
    start.set(qn("w:name"), f"ref_{number}")
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(1000 + number))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def append_citation_hyperlink(paragraph: Paragraph, text: str, ref_number: int) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), f"ref_{ref_number}")
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(vert)
    rpr.append(color)
    run.append(rpr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def first_citation_number(marker: str) -> int:
    match = re.search(r"\d+", marker)
    if not match:
        raise ValueError(marker)
    return int(match.group(0))


def convert_paragraph_citations(paragraph: Paragraph) -> None:
    text = paragraph.text
    if not CITATION_RE.search(text):
        return

    parts: list[tuple[str, str]] = []
    pos = 0
    for match in CITATION_RE.finditer(text):
        if match.start() > pos:
            parts.append(("text", text[pos : match.start()]))
        parts.append(("cite", match.group(0)))
        pos = match.end()
    if pos < len(text):
        parts.append(("text", text[pos:]))

    paragraph.clear()
    for kind, value in parts:
        if kind == "text":
            paragraph.add_run(value)
        else:
            append_citation_hyperlink(paragraph, value, first_citation_number(value))


def replace_reference_list(document: Document) -> None:
    ref_heading = find_paragraph(document, "参考文献")
    try:
        ack_heading = find_paragraph(document, "致  谢")
    except RuntimeError:
        ack_heading = find_paragraph(document, "致 谢")
    template_element = ref_heading._p.getnext()
    if template_element is None:
        raise RuntimeError("参考文献标题后缺少模板段落")
    template = Paragraph(template_element, ref_heading._parent)

    remove_between(ref_heading, ack_heading)
    anchor = ref_heading
    for number, ref in enumerate(FINAL_REFS, 1):
        paragraph = clone_after(anchor, f"[{number}] {ref}", template)
        add_bookmark(paragraph, number)
        anchor = paragraph


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX_PATH)
    args = parser.parse_args()
    docx_path = args.docx

    document = Document(docx_path)

    replace_section(
        find_paragraph(document, "1.3.1 国外研究现状"),
        find_paragraph(document, "1.3.2 国内研究现状"),
        FOREIGN_PARAGRAPHS,
    )
    replace_section(
        find_paragraph(document, "1.3.2 国内研究现状"),
        find_paragraph(document, "1.3.3 现状综述"),
        DOMESTIC_PARAGRAPHS,
    )
    replace_section(
        find_paragraph(document, "1.3.3 现状综述"),
        find_paragraph(document, "1.4 本文的主要工作与论文结构"),
        SUMMARY_PARAGRAPHS,
    )

    for paragraph in document.paragraphs:
        text = paragraph.text
        for old, new in TEXT_REPLACEMENTS.items():
            text = text.replace(old, new)
        if text != paragraph.text:
            clear_and_set(paragraph, text)

    replace_reference_list(document)

    for paragraph in document.paragraphs:
        if ptext(paragraph) == "致 谢":
            clear_and_set(paragraph, "致  谢")

    ref_heading_index = next(i for i, p in enumerate(document.paragraphs) if ptext(p) == "参考文献")
    for paragraph in document.paragraphs[:ref_heading_index]:
        convert_paragraph_citations(paragraph)

    document.save(docx_path)
    print(docx_path)


if __name__ == "__main__":
    main()
