from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Users/sheng/Documents/code/hiking")
SRC = ROOT / "DOCS/thesis/manuscripts/final/陈胜论文.docx"
OUT = ROOT / "DOCS/thesis/manuscripts/revisions/陈胜论文-按批注修订V1.docx"


def set_run_fonts(run, east=None, ascii_font=None, size=None, bold=None):
    if ascii_font:
        run.font.name = ascii_font
    if size:
        run.font.size = size
    if bold is not None:
        run.bold = bold

    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    if east:
        rFonts.set(qn("w:eastAsia"), east)
    if ascii_font:
        rFonts.set(qn("w:ascii"), ascii_font)
        rFonts.set(qn("w:hAnsi"), ascii_font)


def set_para_exact_22(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(22)


def set_first_line_chars_2(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.ind
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if qn("w:firstLine") in ind.attrib:
        del ind.attrib[qn("w:firstLine")]
    ind.set(qn("w:firstLineChars"), "200")


def set_numbering_like(source_para, target_para):
    src_pPr = source_para._p.get_or_add_pPr()
    src_numPr = src_pPr.find(qn("w:numPr"))
    if src_numPr is None:
        return
    tgt_pPr = target_para._p.get_or_add_pPr()
    old = tgt_pPr.find(qn("w:numPr"))
    if old is not None:
        tgt_pPr.remove(old)
    tgt_pPr.append(deepcopy(src_numPr))


def set_cjk_font_in_paragraph(paragraph, east_font, ascii_font="Times New Roman"):
    for run in paragraph.runs:
        if any("\u4e00" <= ch <= "\u9fff" for ch in run.text):
            set_run_fonts(run, east=east_font, ascii_font=ascii_font)


def main():
    OUT.write_bytes(SRC.read_bytes())
    doc = Document(str(OUT))

    # 1) 中文/英文关键词：按批注修复数量、对应关系、分隔形式和局部字体
    doc.paragraphs[36].clear()
    r0 = doc.paragraphs[36].add_run("关键词：")
    set_run_fonts(r0, east="黑体", ascii_font="Times New Roman", size=Pt(12), bold=True)
    r1 = doc.paragraphs[36].add_run("户外路线推荐；GIS；SpringBoot；DeepSeek")
    set_run_fonts(r1, east="仿宋", ascii_font="Times New Roman", size=Pt(12), bold=False)

    doc.paragraphs[41].clear()
    r0 = doc.paragraphs[41].add_run("KEYWORDS: ")
    set_run_fonts(r0, east="宋体", ascii_font="Times New Roman", size=Pt(12), bold=True)
    r1 = doc.paragraphs[41].add_run("Outdoor Route Recommendation; GIS; SpringBoot; DeepSeek")
    set_run_fonts(r1, east="宋体", ascii_font="Times New Roman", size=Pt(12), bold=False)

    # 2) 目录：按批注统一行间距
    for i in range(45, 124):
        set_para_exact_22(doc.paragraphs[i])

    # 2.1 目录里被点名的局部中文字体
    for i in [46, 47, 48, 49]:
        set_cjk_font_in_paragraph(doc.paragraphs[i], "宋体")
    for i in [74, 113, 114, 115, 116, 117, 118, 119]:
        set_cjk_font_in_paragraph(doc.paragraphs[i], "黑体")

    # 3) 正文中被点名的标题/图表题注字体
    set_cjk_font_in_paragraph(doc.paragraphs[243], "黑体")  # 4.1.2 后端服务层
    for i in [536, 539, 545, 548, 551]:
        set_cjk_font_in_paragraph(doc.paragraphs[i], "黑体")  # 6.2.x 测试小节
    for i in [226, 363, 561]:
        set_cjk_font_in_paragraph(doc.paragraphs[i], "宋体")  # 被点名的图/表题注

    # 4) 正文错别字
    doc.paragraphs[310].text = doc.paragraphs[310].text.replace("须求", "需求")

    # 5) 正文段落：按批注统一首行缩进为 2 字符，仅处理正文中的普通段落
    for i in range(125, 577):
        p = doc.paragraphs[i]
        txt = p.text.strip()
        if not txt:
            continue
        if p.style.name not in ("Normal", "Normal (Web)"):
            continue
        if len(txt) <= 10 and txt.endswith("："):
            continue
        set_first_line_chars_2(p)

    # 6) 参考文献：左对齐
    for i in range(578, 606):
        doc.paragraphs[i].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 6.1 最后 4 条参考文献补上与前文一致的编号列表属性
    for i in range(602, 606):
        set_numbering_like(doc.paragraphs[578], doc.paragraphs[i])

    # 6.2 按批注修复 3 条会议文献著录缺少出版地/出版者格式
    doc.paragraphs[580].runs[0].text = (
        "Liu X. Design of personalized tourism route recommendation system based on knowledge graph[C]//"
        "2022 International Conference on Intelligent Transportation, Big Data & Smart City. "
        "[S.l.]: IEEE, 2022: 102-106."
    )
    doc.paragraphs[587].runs[0].text = (
        "Wang Z, Zeng X, Liu W, et al. ToolFlow: Boosting LLM Tool-Calling Through Natural and Coherent "
        "Dialogue Synthesis[C]//Proceedings of NAACL 2025 Long Papers. [S.l.]: ACL, 2025: 4246-4263."
    )
    doc.paragraphs[588].runs[0].text = (
        "Qin S, Zhu Y, Mu L, et al. Meta-Tool: Unleash Open-World Function Calling Capabilities of "
        "General-Purpose Large Language Models[C]//Proceedings of the 63rd Annual Meeting of the "
        "Association for Computational Linguistics. [S.l.]: ACL, 2025: 30653-30677."
    )

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
